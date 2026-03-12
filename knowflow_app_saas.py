#!/usr/bin/env python3
"""
KnowFlow MVP v1.0
YouTube Channel → AI Learning App + High-Ticket Sales Bot
Run: python app.py
"""

import sys, os, re, json, time, threading, tempfile, shutil, subprocess, hashlib
import requests
from pathlib import Path
from datetime import datetime
from http.server import HTTPServer, BaseHTTPRequestHandler
from urllib.parse import urlparse, parse_qs
import webbrowser

# Load .env file
try:
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent / ".env")
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install",
                           "--break-system-packages", "--quiet", "python-dotenv"])
    from dotenv import load_dotenv
    load_dotenv(Path(__file__).parent / ".env")

# Import DB and Auth layers
sys.path.insert(0, str(Path(__file__).parent))
from db import (init_db, get_creator, save_creator, get_all_creators,
                get_knowledge, save_knowledge, append_knowledge, get_word_count,
                get_session, save_session, get_leads, update_subscription,
                update_subscription_by_customer, get_creator_by_clerk_id)
from auth import (verify_clerk_token, extract_token_from_headers, get_clerk_user,
                  create_checkout_session, create_billing_portal_session,
                  handle_stripe_webhook, handle_clerk_webhook, PLANS,
                  CLERK_PUBLISHABLE_KEY)

# ── Auto-install deps ─────────────────────────────────────────────────────────
def pip_install(*packages):
    subprocess.check_call([sys.executable, "-m", "pip", "install",
                           "--break-system-packages", "--quiet", *packages])

def check_deps():
    needs = []
    try: import yt_dlp
    except ImportError: needs.append("yt-dlp")
    try: import youtube_transcript_api
    except ImportError: needs.append("youtube-transcript-api")
    try: import anthropic
    except ImportError: needs.append("anthropic")
    if needs:
        print(f"📦 Installing: {', '.join(needs)}")
        pip_install(*needs)

check_deps()

import yt_dlp
from youtube_transcript_api import YouTubeTranscriptApi
import anthropic

# ── Data storage — now using Postgres ─────────────────────────────────────────
# Keep DATA_DIR for transcript files (local cache during transcription)
DATA_DIR = Path(os.environ.get("DATA_DIR", "/app/data/knowflow_data"))
DATA_DIR.mkdir(parents=True, exist_ok=True)

def creator_dir(slug: str) -> Path:
    d = DATA_DIR / slug
    d.mkdir(exist_ok=True)
    return d

def load_json(path: Path, default):
    if path.exists():
        try: return json.loads(path.read_text())
        except: pass
    return default

def save_json(path: Path, data):
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2))

# Initialize database on startup
try:
    init_db()
except Exception as e:
    print(f"⚠️ DB init warning: {e}")

def parse_multipart(rfile, content_type: str, content_length: int) -> dict:
    """Robust multipart/form-data parser — replaces deprecated cgi module."""
    import re as _re
    bm = _re.search(r'boundary=([^\s;]+)', content_type)
    if not bm:
        return {}
    boundary = bm.group(1).strip('"').encode()
    raw = rfile.read(content_length) if content_length > 0 else rfile.read()
    fields = {}
    delimiter = b'--' + boundary
    for part in raw.split(delimiter)[1:]:
        if part.startswith(b'--') or not part.strip():
            continue
        sep = b'\r\n\r\n' if b'\r\n\r\n' in part else b'\n\n'
        if sep not in part:
            continue
        hdr_block, body = part.split(sep, 1)
        if body.endswith(b'\r\n'):
            body = body[:-2]
        hdrs = hdr_block.decode('utf-8', errors='ignore')
        nm = _re.search(r'name="([^"]+)"', hdrs)
        if not nm:
            continue
        name = nm.group(1)
        fn = _re.search(r'filename="([^"]*)"', hdrs)
        if fn:
            fields[name] = {'filename': fn.group(1), 'data': body}
        else:
            fields[name] = body.decode('utf-8', errors='ignore')
    return fields


# ── YouTube Transcriber v5 (from transcriber.py) ─────────────────────────────

# Whisper pool — one model instance per thread for true parallelism
_whisper_pool    = {}
_whisper_lock    = threading.Lock()
_whisper_ready   = False
_whisper_model_name = "tiny"

def _get_whisper_model():
    tid = threading.get_ident()
    if tid not in _whisper_pool:
        try:
            from faster_whisper import WhisperModel
            _whisper_pool[tid] = WhisperModel(
                _whisper_model_name, device="cpu",
                compute_type="int8", num_workers=1, cpu_threads=2
            )
        except Exception:
            _whisper_pool[tid] = None
    return _whisper_pool.get(tid)

def whisper_transcribe_file(audio_path: str):
    try:
        from faster_whisper import WhisperModel
        model = _get_whisper_model()
        if not model: return None
        segments, _ = model.transcribe(audio_path, beam_size=1, vad_filter=True, word_timestamps=False)
        return " ".join(s.text.strip() for s in segments).strip() or None
    except ImportError: pass
    except Exception: return None
    try:
        import whisper as ow
        if not hasattr(whisper_transcribe_file, "_ow"):
            whisper_transcribe_file._ow = ow.load_model(_whisper_model_name)
        with _whisper_lock:
            result = whisper_transcribe_file._ow.transcribe(audio_path, fp16=False, verbose=False)
        return result.get("text","").strip() or None
    except Exception: return None

def download_audio(video_id: str, tmp_dir: Path):
    ydl_opts = {
        "quiet": True, "no_warnings": True,
        "format": "worstaudio/bestaudio[abr<=64]/bestaudio",
        "outtmpl": str(tmp_dir / f"{video_id}.%(ext)s"),
        "postprocessors": [{"key":"FFmpegExtractAudio","preferredcodec":"mp3","preferredquality":"48"}],
        "socket_timeout": 30, "retries": 2,
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.download([f"https://youtube.com/watch?v={video_id}"])
        files = list(tmp_dir.glob(f"{video_id}.*"))
        return str(files[0]) if files else None
    except: return None

def fetch_yt_transcript(video_id: str):
    """Phase 1: fast YouTube subtitle fetch, no download."""
    try:
        api = YouTubeTranscriptApi()
        tlist = api.list(video_id)
        transcript = None
        for lang in ["de","en","de-DE","en-US","en-GB"]:
            try: transcript = tlist.find_manually_created_transcript([lang]); break
            except: pass
        if not transcript:
            for lang in ["de","en","de-DE","en-US","en-GB"]:
                try: transcript = tlist.find_generated_transcript([lang]); break
                except: pass
        if not transcript:
            all_t = list(tlist)
            if all_t: transcript = all_t[0]
        if not transcript: return None
        parts = []
        for s in transcript.fetch():
            t = s.text if hasattr(s,"text") else s.get("text","")
            t = re.sub(r"<[^>]+>","",t).replace("\n"," ").strip()
            if t: parts.append(t)
        return " ".join(parts) or None
    except: return None

def get_channel_videos(channel_url: str, max_videos: int = 0, include_shorts: bool = True):
    """Scan /videos AND /shorts separately — so longform videos are always found."""
    tlog("🔍 Scanne Channel...", "info")
    base = channel_url.strip().rstrip("/")
    for sfx in ["/videos","/shorts","/streams","/playlists"]:
        if base.endswith(sfx): base = base[:-len(sfx)]; break

    ydl_opts = {"quiet":True,"no_warnings":True,"extract_flat":"in_playlist","ignoreerrors":True}
    all_videos = {}
    channel_name = "Unbekannt"

    scan_urls = [base + "/videos"]
    if include_shorts:
        scan_urls.append(base + "/shorts")

    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        for url in scan_urls:
            section = "Shorts" if url.endswith("/shorts") else "Videos"
            try:
                info = ydl.extract_info(url, download=False)
                if not info: continue
                entries = []
                for e in (info.get("entries") or []):
                    if not e: continue
                    if "entries" in e:
                        entries.extend([x for x in e["entries"] if x and x.get("id")])
                    elif e.get("id"):
                        entries.append(e)
                if entries and channel_name == "Unbekannt":
                    channel_name = (info.get("channel") or info.get("uploader") or info.get("title") or "Unbekannt")
                    tlog(f"✅ Channel: <b>{channel_name}</b>", "success")
                for e in entries:
                    vid_id = e.get("id")
                    if not vid_id or vid_id in all_videos: continue
                    dur = e.get("duration") or 0
                    all_videos[vid_id] = {
                        "id": vid_id,
                        "title": e.get("title","Ohne Titel"),
                        "duration": dur,
                        "is_short": url.endswith("/shorts") or (0 < dur <= 65),
                        "upload_date": e.get("upload_date",""),
                    }
                tlog(f"📋 {section}: <b>{len(entries)}</b> gefunden", "info")
            except Exception as ex:
                tlog(f"⚠️ Fehler bei {section}: {ex}", "warn")

    videos = list(all_videos.values())
    longform = [v for v in videos if not v["is_short"]]
    shorts   = [v for v in videos if v["is_short"]]
    tlog(f"🎬 Longform: <b>{len(longform)}</b> | Shorts: <b>{len(shorts)}</b>", "info")
    result = longform + (shorts if include_shorts else [])
    if max_videos > 0:
        result = result[:max_videos]
        tlog(f"⚙️ Limitiert auf <b>{max_videos}</b>", "warn")
    return result, channel_name

# ── Global transcription progress ────────────────────────────────────────────
transcribe_log = []
transcribe_done = False
transcribe_slug = None
log_lock = threading.Lock()

def tlog(msg, t="info"):
    with log_lock:
        transcribe_log.append({"msg": msg, "type": t, "ts": time.time()})

def ensure_faster_whisper():
    """Install and verify faster-whisper, fallback to openai-whisper."""
    global _whisper_ready
    try:
        from faster_whisper import WhisperModel
        _whisper_ready = True
        tlog("✅ faster-whisper bereit", "success")
        return True
    except ImportError:
        tlog("📦 Installiere faster-whisper...", "info")
        try:
            subprocess.check_call([sys.executable, "-m", "pip", "install",
                "--break-system-packages", "--quiet", "faster-whisper"])
            from faster_whisper import WhisperModel
            _whisper_ready = True
            tlog("✅ faster-whisper installiert", "success")
            return True
        except Exception as e:
            tlog(f"⚠️ faster-whisper fehlgeschlagen, versuche openai-whisper...", "warn")
            try:
                subprocess.check_call([sys.executable, "-m", "pip", "install",
                    "--break-system-packages", "--quiet", "openai-whisper"])
                _whisper_ready = True
                tlog("✅ openai-whisper als Fallback bereit", "success")
                return True
            except Exception as e2:
                tlog(f"❌ Kein Whisper verfügbar: {e2}", "error")
                return False

def assemblyai_transcribe(video_id: str) -> str:
    """Transcribe a YouTube video via AssemblyAI using youtube-transcript-api with different language attempts."""
    # First try extended transcript API with more language options
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        api = YouTubeTranscriptApi()
        # Try all available transcripts including auto-generated
        transcript_list = api.list(video_id)
        for t in transcript_list:
            try:
                snippets = t.fetch()
                text = " ".join(s.text for s in snippets).strip()
                if text and len(text) > 100:
                    return text
            except:
                continue
    except Exception:
        pass

    # Fallback: AssemblyAI with direct YouTube URL (works for public videos)
    api_key = os.environ.get("ASSEMBLYAI_API_KEY", "")
    if not api_key:
        return None
    try:
        headers = {"authorization": api_key, "content-type": "application/json"}
        youtube_url = f"https://www.youtube.com/watch?v={video_id}"
        resp = requests.post(
            "https://api.assemblyai.com/v2/transcript",
            headers=headers,
            json={"audio_url": youtube_url, "language_detection": True},
            timeout=30
        )
        if resp.status_code != 200:
            print(f"AssemblyAI submit failed for {video_id}: {resp.status_code} {resp.text[:200]}")
            return None
        job_id = resp.json().get("id")
        if not job_id:
            return None
        for _ in range(120):
            time.sleep(5)
            poll = requests.get(
                f"https://api.assemblyai.com/v2/transcript/{job_id}",
                headers=headers, timeout=30
            )
            status = poll.json().get("status")
            if status == "completed":
                return poll.json().get("text", "").strip() or None
            elif status == "error":
                print(f"AssemblyAI error for {video_id}: {poll.json().get('error')}")
                return None
        return None
    except Exception as e:
        print(f"AssemblyAI error for {video_id}: {e}")
        return None

def phase2_assemblyai(videos_no_yt: list, workers: int = 5) -> dict:
    """Phase 2: AssemblyAI transcription for videos without YT subtitles."""
    results = {}
    lock = threading.Lock()
    done_c = [0]
    total = len(videos_no_yt)
    tlog(f"🎙️ <b>Phase 2:</b> {total} Videos per AssemblyAI ({workers} parallel)...", "info")

    def worker(video):
        vid_id = video["id"]
        text = assemblyai_transcribe(vid_id)
        with lock:
            results[vid_id] = text
            done_c[0] += 1
            cnt = done_c[0]
        short_t = video["title"][:45] + "…" if len(video["title"]) > 45 else video["title"]
        if text:
            tlog(f"🎙️ {cnt}/{total} — {short_t} — <b>{len(text.split()):,} Wörter</b>", "success")
        else:
            tlog(f"⚠️ {cnt}/{total} — {short_t} — fehlgeschlagen", "warn")

    from concurrent.futures import ThreadPoolExecutor, as_completed
    with ThreadPoolExecutor(max_workers=workers) as ex:
        list(as_completed([ex.submit(worker, v) for v in videos_no_yt]))

    ok = sum(1 for v in results.values() if v)
    tlog(f"✅ Phase 2 fertig: <b>{ok}/{total}</b> per AssemblyAI", "success")
    return results

def run_transcription_bg(slug: str, channel_url: str, max_videos: int):
    global transcribe_done, transcribe_slug, _whisper_pool, _whisper_ready, _whisper_model_name
    transcribe_slug = slug
    transcribe_done = False
    transcribe_log.clear()
    _whisper_pool = {}
    _whisper_ready = False

    cdir = creator_dir(slug)
    config_path = cdir / "config.json"
    transcript_path = cdir / "transcript.txt"

    try:
        videos, channel_name = get_channel_videos(channel_url, max_videos, include_shorts=False)
        total = len(videos)
        if not videos:
            tlog("❌ Keine Videos gefunden", "error")
            return

        config = load_json(config_path, {})
        config.update({"channel_name": channel_name, "channel_url": channel_url,
                        "video_count": total, "last_updated": datetime.now().isoformat()})
        save_json(config_path, config)

        # ── Phase 1: YouTube subtitles (parallel, fast, no download) ─────────
        tlog(f"⚡ <b>Phase 1:</b> {total} Videos — YouTube-Untertitel (16 parallel)...", "info")
        yt_results = {}
        lock = threading.Lock()
        done_count = [0]

        def fetch_worker(video):
            text = fetch_yt_transcript(video["id"])
            with lock:
                yt_results[video["id"]] = text
                done_count[0] += 1
                cnt = done_count[0]
            short_t = video["title"][:45] + "…" if len(video["title"]) > 45 else video["title"]
            if text:
                tlog(f"✅ {cnt}/{total} — {short_t} — <b>{len(text.split()):,} Wörter</b>", "success")
            else:
                tlog(f"⬜ {cnt}/{total} — {short_t} — kein Untertitel", "dim")

        from concurrent.futures import ThreadPoolExecutor, as_completed
        with ThreadPoolExecutor(max_workers=16) as ex:
            list(as_completed([ex.submit(fetch_worker, v) for v in videos]))

        no_yt = [v for v in videos if not yt_results.get(v["id"])]
        got_yt = total - len(no_yt)
        tlog(f"📊 Phase 1: <b>{got_yt} mit Untertiteln</b> | <b>{len(no_yt)} ohne</b>", "info")

        # ── Phase 2: AssemblyAI für Videos ohne Untertitel ───────────────────
        whisper_results = {}
        if no_yt:
            assembly_key = os.environ.get("ASSEMBLYAI_API_KEY", "")
            if assembly_key:
                tlog(f"🤖 Starte AssemblyAI für {len(no_yt)} Videos...", "info")
                whisper_results = phase2_assemblyai(no_yt, workers=5)
            else:
                tlog(f"⚠️ ASSEMBLYAI_API_KEY nicht gesetzt — {len(no_yt)} Videos ohne Transkript", "warn")

        # ── Write transcript file ─────────────────────────────────────────────
        all_text = []
        with open(transcript_path, "w", encoding="utf-8") as f:
            f.write(f"KNOWFLOW KNOWLEDGE BASE\nChannel: {channel_name}\n")
            f.write(f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
            f.write(f"Videos: {total}\n{'=' * 60}\n\n")
            for video in videos:
                vid_id = video["id"]
                text = yt_results.get(vid_id) or whisper_results.get(vid_id)
                method = "YouTube" if yt_results.get(vid_id) else "Whisper" if whisper_results.get(vid_id) else None
                vtype = "Short" if video.get("is_short") else "Longform"
                if text:
                    f.write(f"### {video['title']} [{vtype}] [{method}]\n")
                    f.write(f"URL: https://youtube.com/watch?v={vid_id}\n\n")
                    f.write(text + "\n\n")
                    all_text.append(text)
                else:
                    f.write(f"--- KEIN TRANSKRIPT: {video['title']} [{vtype}]\n\n")

        yt_ok = sum(1 for v in videos if yt_results.get(v["id"]))
        wh_ok = sum(1 for v in videos if whisper_results.get(v["id"]))
        total_words = sum(len(t.split()) for t in all_text)
        tlog(f"🎉 <b>Fertig! {len(all_text)}/{total} Transkripte — {total_words:,} Wörter</b>", "done")
        tlog(f"📊 YouTube: <b>{yt_ok}</b> | Whisper: <b>{wh_ok}</b> | Keine: <b>{total - yt_ok - wh_ok}</b>", "info")
        tlog(f"📄 Gespeichert: {transcript_path}", "file")

        # ── Save to DB ────────────────────────────────────────────────────────
        try:
            full_text = transcript_path.read_text(encoding="utf-8")
            # Ensure creator exists in DB before saving knowledge
            existing = get_creator(slug)
            if not existing:
                save_creator(slug, {
                    "channel_name": channel_name,
                    "channel_url": channel_url,
                    "video_count": total,
                })
            else:
                save_creator(slug, {
                    "channel_name": channel_name,
                    "channel_url": channel_url,
                    "video_count": total,
                })
            save_knowledge(slug, full_text)
            tlog(f"✅ In Datenbank gespeichert", "success")
        except Exception as e:
            tlog(f"⚠️ DB-Speicherung fehlgeschlagen: {e}", "warn")

    except Exception as e:
        import traceback
        tlog(f"❌ Fehler: {e}", "error")
        tlog(traceback.format_exc()[-500:], "error")
    finally:
        transcribe_done = True

# ── AI Engine ─────────────────────────────────────────────────────────────────
def load_knowledge(slug: str) -> str:
    # First try DB
    knowledge = get_knowledge(slug)
    if knowledge:
        return knowledge[:80000]
    # Fallback: local transcript file (for existing data)
    path = creator_dir(slug) / "transcript.txt"
    if path.exists():
        text = path.read_text(encoding="utf-8")
        # Migrate to DB on first access
        try:
            save_knowledge(slug, text)
            print(f"✅ Migriert: {slug} transcript → DB")
        except Exception as e:
            print(f"⚠️ Migration failed: {e}")
        return text[:80000]
    return ""    

SALES_SYSTEM_PROMPT = """You are an elite AI knowledge assistant and world-class high-ticket sales closer for a creator's KnowFlow page.

You have two roles you blend naturally and seamlessly:

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ROLE 1: KNOWLEDGE EXPERT
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Answer questions about the creator's content using ONLY their transcript knowledge base.
- Be thorough, specific, reference actual content from their videos
- Be genuinely helpful and educational
- Generate quiz questions when asked
- Never invent content that isn't in the knowledge base

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
ROLE 2: DYNAMIC HIGH-TICKET SALES CLOSER
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

## THE CORE PRINCIPLE — NOT A SCRIPT, A DIAGNOSTIC SYSTEM

You are NOT running a linear script. You are a diagnostician who asks after every message:
"What is MISSING right now for this person to make a decision?"

Four psychological conditions must ALL be present before you present a solution:
  1. DESIRE         → They vividly see and feel the life they want
  2. URGENCY        → They feel the real, quantified COST of staying stuck (NEED, not just pain)
  3. BELIEF         → They believe the solution works AND they are capable of doing it
  4. RESPONSIBILITY → They own the decision (not outsourced to money/partner/time)

If any is missing, return to the block that creates it. Never push forward prematurely.

---

## THE PRE-FRAME METHOD — YOUR MOST IMPORTANT WEAPON

Objections do NOT appear at the end of the call. They exist BEFORE the call starts.
Every objection is a BELIEF formed from a past EXPERIENCE.
Your job: find the root experience → reframe the belief → the objection disappears.

EARLY in discovery, always ask: "What have you done in the past to solve this? Why didn't you act sooner?"
→ They give a surface excuse (money, time) → DO NOT wait until end to handle this. Reframe it NOW:
  - Money excuse: "If you had the skills we're talking about, would you have more money? So is the real issue no money — or not yet having the right skills to generate more?"
  - Time excuse: "There are two types of people — those who say 'I don't have time' and those who ask 'how do I find the time?' Which person gets to where they want to go?"
→ Get their commitment to overcome that belief NOW → congruency principle: they cannot raise it at the end without contradicting themselves.
→ Then consequence it: "So if you continue letting [that belief] guide your decisions, where are you in a year?"
→ Then anchor: "So are you committed to letting go of that way of thinking?"

---

## PAIN vs. NEED — THE DIFFERENCE BETWEEN 25% AND 60% CLOSE RATE

PAIN alone does NOT create action. Pain creates complaint.
NEED creates urgency. Your job is to convert their pain into a quantified, tangible NEED.

Present Tension questions:
- "What is this costing you per month, right now, in real money?"
- "What are you losing every month you stay in this position?"
- "If nothing changes, what are the chances you achieve [goal] in the next 90 days?"

Then quantify: "So you've left roughly €[X] on the table over [Y] months. Does that feel real to you?"

Identify their Maslow level to speak the right language:
- Level 2 (Security): speak to financial stability, avoiding loss, consistency
- Level 3 (Belonging): speak to community, not being alone in this, being understood
- Level 4 (Esteem): speak to status, being respected, being seen as capable and elite
- Level 5 (Self-actualization): speak to legacy, purpose, becoming who they know they can be

---

## THE FAVORABLE CAUSE — CRITICAL AND USUALLY MISSED

Two separate concepts:
1. PERCEIVED PROBLEM: What they feel is stopping them (not enough leads, stuck, frustrated)
2. FAVORABLE CAUSE: WHY the problem exists = the thing YOU solve

If they believe they can fix the cause themselves → "I'll figure it out on my own" → no sale.
Your job: Build doubt that they can solve it alone using the 3-part Doubt Sequence:

  Step A — Accept the cause (as a question, never a statement):
    "Have you considered that the reason you're getting [problem] is because you don't have [favorable cause]? Does that make sense?"
  Step B — Confusion question (should stump them):
    "In your current process, what do you have in place to [solve the cause in a detailed, specific way]?"
    → The answer you want: "I have nothing for that."
    → This builds doubt that they can fix it alone.
  Step C — Link it together:
    "Do you see why that's causing [problem]? Does that make sense?"

Once they have DOUBT they can solve it alone → they are in the buying pocket.

---

## DRAWN-OUT IMPACTS — BUILDING PROBLEM DEPTH

Every problem has a cascade. Surface each level through questions, never through statements.
Map the cascade for this offer, then ask questions that lead there naturally.

Example cascade (adapt to creator's context):
  Weak content → No leads → No calls → No income → Stuck → Family/life goals unreachable

For EACH impact, use contextualization so the question doesn't feel obvious or pushy:
  "I know some people in this position don't really mind X — for you, how is [problem] affecting [next impact]?"

Always end the cascade with an EMOTIONAL consequence:
  "And personally — how does it feel knowing that this has been going on for [X months]?"
  "Are you willing to settle for that as your reality?"

The bigger the problem they feel, the more natural and inevitable the solution becomes.

---

## IDENTITY SELLING — THE UPGRADE FROM OUTCOME SELLING

Don't just sell outcomes ("you'll make more money"). Sell IDENTITY.

Help them see themselves as the person who has already achieved it:
- "When you get to that point — who does that make you? How does that feel?"
- "What does it mean for you PERSONALLY — not just financially — to be that person?"
- "Being the person who can [outcome] — how do you see yourself differently?"
- "If you fast-forward one year and you're that person — what has changed in your day-to-day?"

Then create contrast: "Compared to how you feel NOW — what's the gap between those two people?"

Two beliefs must be in place before the sale closes:
  1. They believe in THEMSELVES — they won't fail if they commit
  2. They believe in the CREATOR/OFFER — it will deliver for them specifically

If either is missing, address it before moving forward.

---

## THE 10 DYNAMIC BLOCKS

### BLOCK 1 — FRAME
Set authority and expectations at the start.
"My job is to understand where you're at, where you want to go, and whether this can help.
If it can, I'll walk you through it. If not, I'll be honest. Fair enough?"

### BLOCK 2 — ALIGNMENT
Find their goal and real motivation behind the goal.
"What made you look into this further?"
"If everything worked exactly as you wanted — what would your situation look like?"
"Why is that goal important to you specifically?"

### BLOCK 3 — PROBLEM DISCOVERY (5 Categories)
Current state → Desired state → Gap analysis → Root cause → Limiting beliefs
"Walk me through what you've been doing up until now."
"How long has that been going on?"
"What have you tried so far? Why didn't that work?"

### BLOCK 4 — FAVORABLE CAUSE + DOUBT BUILDING
→ Use the 3-part Doubt Sequence.
→ Get them to agree: the reason they can't fix it alone is exactly what you solve.

### BLOCK 5 — PRE-FRAME (Early Objection Reframe)
→ "What's actually held you back from solving this sooner?"
→ Surface the excuse early, reframe the belief, get a commitment to overcome it.
→ Now they cannot raise it as an objection at the end (congruency).

### BLOCK 6 — DRAWN-OUT IMPACTS
→ Cascade problems through questions with contextualization.
→ Quantify the loss in real numbers.
→ End with emotional impact + "Are you willing to settle for that?"

### BLOCK 7 — IDENTITY + FUTURE VISION
→ Paint the desired identity in vivid emotional detail.
→ Get them to FEEL what it's like to be the person who has achieved it.
→ Create contrast with how they feel now.
→ GAP: "With how things are going now, how close are you to being that person?"
→ Cost of inaction: "What if you never get to experience being that person because [current pattern]?"

### BLOCK 8 — SOLUTION PRESENTATION
Only present when all 4 conditions are confirmed.
Use THEIR exact words to connect problems to pillars.
"Based on what you shared — the issue isn't effort. It's that you haven't had [X]."
3-Pillar structure: each pillar = one of THEIR problems → your solution → THEIR desired outcome.
Check: "On a scale of 1-10 — where do you find yourself on this?"

### BLOCK 9 — CLOSE
Decision extract BEFORE price: "Do you feel like this is the right direction for you?"
Then: "So the question is — do you keep figuring this out alone, or do you want help getting there faster?"
Then and ONLY then: booking link or payment options.

### BLOCK 10 — OBJECTION HANDLING (Precise Live-Call Patterns)

CRITICAL: Surface objections are SYMPTOMS, never the real issue. Diagnose first.

**MONEY:**
First check belief: "Aside from the investment — do you feel this can actually get you to [desired state]?"
  → YES (logistics/nerves): "Is it more about nerves, or do we need to structure a payment plan?"
    • Reality Strategy: "Has there ever been a time you didn't have money but found a way because it mattered enough to you?"
    • Reframe: "It's not that you don't have the money — it's that you don't yet have the skills to generate more. The investment IS the solution to the money problem."
    • Hierarchy: "Isn't it more important to take a calculated risk now than to keep losing €X every month?"
    • If still stuck: confirm total available → structure split payment → close
  → NO (belief problem): Return to Block 4+7. Build Favorable Cause + Identity, then re-close.

**PARTNER/SPOUSE:**
Pre-handle early: "When it comes to investing in yourself — is that a decision you make, or do you and your partner decide together?"
  → "I decide" → "So if we both feel this is the right fit today, you'd move forward?"
Late objection — diagnose: logistical (shared accounts, equity partner) OR fear?
  → Test: "If your partner was right here and fully on board — would YOU do it?"
  → YES → fear underneath → find it and handle it
  → NO → genuine logistics → help them present it to partner + secure a deposit now
  Fear route: "Who is responsible for getting your family where you want them to be?"
  → "Heavy is the head that wears the crown. The burden of leadership falls on the decision maker."
  → Consequence: "If you go to them, they say no, and you stay exactly where you are — what does that mean for you?"
  Trust frame: "She trusts you to invest in yourself and your career, right? So if she trusts you — do you trust yourself?"

**THINK ABOUT IT:**
"Aside from thinking about it — do you feel the program can get you to [desired state]?"
  → YES → "What specifically do you need to think through? Is it the investment, timing, or something else?"
    → Routes to Money / Time / Partner / Self-belief → handle the real one
  → NO → belief not built → return to Block 4+7
  → "Has there ever been a perfect time in your life to start anything?"
  → "The people making [desired outcome] didn't wait for the right time. They made THIS the right time."

**TIME:**
"Aside from time — do you feel the program can get you there?"
  → "There are two types of people: those who say 'I don't have time' and those who ask 'how do I make time.' Which one gets where they want to go?"
  → Reframe: "It's not a time problem. It's a priority problem."
  → "What has actually changed that would make waiting another year different from the last year?"

**FEAR / UNCERTAINTY:**
Master question: "If you knew with 100% certainty it would work — would you do it?"
  → YES → self-belief issue → Identity Selling (Block 7)
    • "The last time you did something uncomfortable — what got better because of it?"
    • LeBron analogy: "Even LeBron has coaches. The most successful people don't do it alone."
    • Valley Metaphor: "Right now you're standing in the valley. From there, all you see is valley. The life you want is at the summit. You don't need to see the whole path — you just need to take the first step."
    • Fat person analogy: "A person who keeps making decisions like someone who's out of shape — where do they end up? The mindset has to change before the results can."
  → NO → skepticism about offer → more social proof, more connection to their specific situation

**LOOPING after objection handling:**
Elite closers loop 3-4 times naturally. Never give up after one loop.
Loop pattern: Objection → diagnose real issue → return to correct block → re-close → repeat if needed.
"If you keep making decisions the same way you have been — where will you be in one year?"

---

## CONGRUENCY PRINCIPLE
What someone says out loud, they find hard to contradict.
→ Get micro-commitments throughout: "So you ARE committed to changing this?"
→ Use identity labels before deep questions: "I respect how self-aware you are about this..."
→ Reference their own words back: "You told me [exact words] — so when you say [objection]..."
→ Make objections feel incongruent FOR THEM, not something you argue against.

---

## TACTICAL VULNERABILITY
Before asking a deep or personal question, briefly share a real story first.
"When I was in a similar situation, what pushed me was [X]. For you — what's really driving you to do something about this now?"
This creates psychological safety and makes them feel obligated to reciprocate openness.

---

## TONE AND STYLE (Non-negotiable)
- **ONE question at a time. NEVER stack multiple questions. End EVERY response with max ONE question.**
- Keep responses concise — 2-4 sentences max, then ONE question. Do not dump everything at once.
- No bullet-point walls. Write like a real conversation, not a presentation.
- Mirror their EXACT language back to them.
- Softeners: "Just so I have context...", "Walk me back through...", "Give me an idea...", "Can I share a perspective on that?"
- You are on THEIR TEAM against their problem. Never adversarial.
- You are a DIAGNOSTICIAN. You surface internal conflict through questions.
- The sale happens when THEY realize staying the same is more painful than changing.
- Never tell — always ask. When you tell, they resist. When they conclude, they buy.
- Blend knowledge answers with sales diagnostic seamlessly. It should never feel like a sales call.
- Format responses cleanly. No markdown asterisks as decoration. Use bold sparingly for truly key phrases only.

---

## WARM LEAD DETECTION
Flag as warm lead when prospect asks about: price, cost, how much, how to start, sign up, booking, ready, next steps, working directly with creator.
→ "The next step is a strategy call where you both confirm it's the right fit. You can book directly here: [BOOKING_LINK]"

---

## CREATOR CONTEXT:
{creator_context}

## PRODUCTS:
{products}

## BOOKING LINK:
{booking_link}

## KNOWLEDGE BASE:
{knowledge_base}

---
Remember: Lead with genuine value. Educate freely. Build problems deeply before presenting solutions. The sale is already made in the discovery — if you do it right, the close is just a formality.
"""

def get_ai_response(slug: str, messages: list, session_data: dict) -> dict:
    config = get_creator(slug)
    knowledge = load_knowledge(slug)

    products_str = ""
    for p in config.get("products", []):
        products_str += f"- {p.get('name', '')}: {p.get('description', '')} → {p.get('url', '')}\n"

    # Build diagnostic state summary for the AI
    msg_count = len(messages)
    diag_state = session_data.get("diag_state", {
        "desire": False, "urgency": False, "belief": False, "responsibility": False,
        "current_block": "ALIGNMENT", "objections_seen": []
    })

    diagnostic_context = f"""
## CURRENT SESSION STATE (message #{msg_count})
Current diagnostic block: {diag_state.get('current_block', 'ALIGNMENT')}
Desire established: {diag_state.get('desire', False)}
Urgency established: {diag_state.get('urgency', False)}
Belief established: {diag_state.get('belief', False)}
Responsibility established: {diag_state.get('responsibility', False)}
Objections seen so far: {', '.join(diag_state.get('objections_seen', [])) or 'none yet'}

Instruction: Based on the conversation so far, identify which of the 4 conditions is currently missing and address it. Do NOT present the solution yet unless all 4 are confirmed.
"""

    system = SALES_SYSTEM_PROMPT.format(
        creator_context=config.get("bio", "A content creator sharing valuable knowledge."),
        products=products_str or "No products configured yet.",
        booking_link=config.get("booking_link", "Not configured yet."),
        knowledge_base=knowledge[:55000] if knowledge else "No transcript loaded yet."
    ) + diagnostic_context + """

WICHTIG — Antwortlänge: Halte Antworten kurz und präzise (max. 3-4 Sätze oder 4-5 Bullet Points).
Der Zuschauer fragt nach wenn er mehr wissen will. Kein Markdown-Fettdruck (**text**), keine Sternchen."""

    client = anthropic.Anthropic()

    # Detect warm lead signals in last user message
    last_user_msg = ""
    for m in reversed(messages):
        if m.get("role") == "user":
            last_user_msg = m.get("content", "").lower()
            break

    # Only trigger warm lead after meaningful conversation depth
    warm_signals = [
        "preis", "kosten", "wie viel kostet", "how much",
        "book", "buchen", "ready", "bereit",
        "kaufen", "invest", "join", "mitmachen",
        "nächster schritt", "next step", "zusammenarbeiten",
        "payment", "zahlung", "how do i work with you", "wie kann ich mit dir"
    ]
    # Require depth: at least 6 messages AND desire+urgency both present
    desire_ok = diag_state.get("desire", False)
    urgency_ok = diag_state.get("urgency", False)
    is_warm = msg_count >= 6 and (desire_ok or urgency_ok) and any(signal in last_user_msg for signal in warm_signals)

    # Detect objection type in last message
    objection_map = {
        "money": ["don't have the money", "can't afford", "too expensive", "kein geld", "zu teuer", "investment"],
        "time": ["don't have time", "too busy", "keine zeit", "zu beschäftigt"],
        "partner": ["ask my partner", "ask my spouse", "ask my wife", "ask my husband", "partner fragen"],
        "think_about_it": ["need to think", "let me think", "think about it", "überlegen"]
    }
    detected_objection = None
    for obj_type, signals in objection_map.items():
        if any(s in last_user_msg for s in signals):
            detected_objection = obj_type
            break

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=800,
            system=system,
            messages=messages
        )
    except Exception as api_err:
        err_str = str(api_err)
        if "401" in err_str or "api-key" in err_str.lower() or "authentication" in err_str.lower():
            raise RuntimeError("⚠️ API Key fehlt oder ungültig. Server neu starten mit: export ANTHROPIC_API_KEY=sk-ant-... && python3 knowflow_app.py")
        raise

    response_text = response.content[0].text

    # Detect warm lead from AI response — only after msg 8 and conditions met
    ai_suggests_call = msg_count >= 8 and any(
        signal in response_text.lower()
        for signal in ["strategy call", "booking link", "buche einen call", "nächsten schritt"]
    )
    warm_lead_detected = is_warm or ai_suggests_call

    # Update diagnostic state based on conversation signals
    full_conv = " ".join(m.get("content", "") for m in messages).lower()
    diag_state["desire"] = any(w in full_conv for w in ["want", "goal", "achieve", "dream", "vision", "ziel", "möchte"])
    diag_state["urgency"] = any(w in full_conv for w in ["how long", "year", "waiting", "still the same", "wie lange", "problem"])
    diag_state["belief"] = msg_count >= 6  # Rough heuristic — by msg 6 belief has been tested
    diag_state["responsibility"] = not any(w in full_conv for w in ["partner", "spouse", "think about", "not sure", "überlegen"])
    if detected_objection and detected_objection not in diag_state.get("objections_seen", []):
        diag_state.setdefault("objections_seen", []).append(detected_objection)

    # Advance block tracking
    if msg_count <= 2:
        diag_state["current_block"] = "FRAME + ALIGNMENT"
    elif msg_count <= 5:
        diag_state["current_block"] = "PROBLEM DISCOVERY"
    elif msg_count <= 8:
        diag_state["current_block"] = "ROOT CAUSE + CONSEQUENCES"
    elif msg_count <= 12:
        diag_state["current_block"] = "FUTURE VISION + GAP"
    elif warm_lead_detected:
        diag_state["current_block"] = "CLOSE"
    else:
        diag_state["current_block"] = "SOLUTION / OBJECTION HANDLING"

    return {
        "response": response_text,
        "warm_lead": warm_lead_detected,
        "objection": detected_objection,
        "diag_state": diag_state,
        "booking_link": config.get("booking_link", ""),
        "products": config.get("products", [])
    }

# ── HTTP Handler ──────────────────────────────────────────────────────────────
class Handler(BaseHTTPRequestHandler):
    def log_message(self, *a): pass

    def send_json(self, data, status=200):
        import datetime
        def default_serializer(obj):
            if isinstance(obj, (datetime.datetime, datetime.date)):
                return obj.isoformat()
            raise TypeError(f'Object of type {obj.__class__.__name__} is not JSON serializable')
        body = json.dumps(data, ensure_ascii=False, default=default_serializer).encode()
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Content-Length", len(body))
        self.end_headers()
        self.wfile.write(body)

    def send_html(self, html: str):
        body = html.encode()
        self.send_response(200)
        self.send_header("Content-Type", "text/html; charset=utf-8")
        self.send_header("Content-Length", len(body))
        self.end_headers()
        self.wfile.write(body)

    def read_body(self):
        length = int(self.headers.get("Content-Length", 0))
        return json.loads(self.rfile.read(length)) if length else {}

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        self.end_headers()

    def do_GET(self):
        p = urlparse(self.path)
        qs = parse_qs(p.query)

        if p.path == "/" or p.path == "/creator":
            self.send_html(CREATOR_HTML)

        elif p.path.startswith("/c/"):
            slug = p.path.split("/c/")[1].split("/")[0]
            config = load_json(creator_dir(slug) / "config.json", {})
            if not config:
                self.send_html("<h1>Creator not found</h1>")
            else:
                self.send_html(VIEWER_HTML.replace("{{SLUG}}", slug)
                               .replace("{{CREATOR_NAME}}", config.get("channel_name", slug))
                               .replace("{{CREATOR_BIO}}", config.get("bio", ""))
                               .replace("{{BOOKING_LINK}}", config.get("booking_link", "#")))

        elif p.path == "/api/transcribe/start":
            slug = qs.get("slug", [""])[0]
            url = qs.get("url", [""])[0]
            max_v = int(qs.get("max", [30])[0])
            if not slug or not url:
                self.send_json({"error": "Missing slug or url"}, 400)
                return
            threading.Thread(target=run_transcription_bg, args=(slug, url, max_v), daemon=True).start()
            self.send_json({"ok": True})

        elif p.path == "/api/transcribe/progress":
            from_idx = int(qs.get("from", [0])[0])
            self.send_json({
                "entries": transcribe_log[from_idx:],
                "done": transcribe_done,
                "slug": transcribe_slug
            })

        elif p.path == "/api/creator/config":
            slug = qs.get("slug", [""])[0]
            config = get_creator(slug)
            config["word_count"] = get_word_count(slug)
            self.send_json(config)

        elif p.path in ("/admin", "/admin/"):
            slug = qs.get("slug", [""])[0]
            if not slug:
                # Redirect to login
                self.send_response(302)
                self.send_header("Location", "/login")
                self.end_headers()
                return
            admin_page = (ADMIN_HTML
                .replace("{{SLUG}}", slug)
                .replace("{{CLERK_PK}}", CLERK_PUBLISHABLE_KEY or ""))
            self.send_html(admin_page)
            return

        elif p.path == "/login":
            page = LOGIN_HTML.replace("{{CLERK_PK}}", CLERK_PUBLISHABLE_KEY or "")
            self.send_html(page)
            return

        elif p.path == "/auth/callback":
            page = CALLBACK_HTML.replace("{{CLERK_PK}}", CLERK_PUBLISHABLE_KEY or "")
            self.send_html(page)
            return

        elif p.path == "/setup":
            # Keep for backward compat — redirect to /admin if slug given
            slug = qs.get("slug", [""])[0]
            if slug:
                self.send_response(302)
                self.send_header("Location", f"/admin?slug={slug}")
                self.end_headers()
            else:
                self.send_response(302)
                self.send_header("Location", "/")
                self.end_headers()
            return
        elif p.path == "/api/creators":
            # Admin only
            token = extract_token_from_headers(self.headers)
            user = verify_clerk_token(token) if token else {}
            admin_email = os.environ.get("ADMIN_EMAIL", "sven.gold.official@gmail.com")
            if user.get("email") != admin_email and not os.environ.get("DEV_MODE"):
                self.send_json({"error": "Unauthorized"}, 403)
                return
            self.send_json(get_all_creators())

        elif p.path == "/api/video-status":
            slug = qs.get("slug", [""])[0]
            cfg = load_json(creator_dir(slug) / "config.json", {})
            video_file = cfg.get("greeting_video", "")
            video_path = creator_dir(slug) / video_file if video_file else None
            # Auto-detect video file if config missing the key
            if not video_file:
                for ext in [".mp4", ".mov", ".webm", ".MP4", ".MOV"]:
                    candidate = creator_dir(slug) / f"greeting{ext}"
                    if candidate.exists():
                        video_file = f"greeting{ext}"
                        cfg["greeting_video"] = video_file
                        save_json(creator_dir(slug) / "config.json", cfg)
                        video_path = candidate
                        break
            exists = video_path.exists() if video_path else False
            self.send_json({"slug": slug, "video_file": video_file, "exists": exists,
                           "config_has_key": bool(cfg.get("greeting_video"))})

        elif p.path.startswith("/api/greeting-video/"):
            slug = p.path.split("/api/greeting-video/")[1].split("?")[0].split("/")[0]
            cfg = load_json(creator_dir(slug) / "config.json", {})
            video_filename = cfg.get("greeting_video", "")
            if not video_filename:
                self.send_response(404); self.end_headers(); return
            video_path = creator_dir(slug) / video_filename
            if not video_path.exists():
                self.send_response(404); self.end_headers(); return
            ext = video_path.suffix.lower()
            mime = {"mp4": "video/mp4", "mov": "video/quicktime", "webm": "video/webm"}.get(ext[1:], "video/mp4")
            data_bytes = video_path.read_bytes()
            self.send_response(200)
            self.send_header("Content-Type", mime)
            self.send_header("Content-Length", str(len(data_bytes)))
            self.send_header("Accept-Ranges", "bytes")
            self.send_header("Cache-Control", "no-cache")
            self.send_header("Access-Control-Allow-Origin", "*")
            self.end_headers()
            self.wfile.write(data_bytes)


        elif p.path == "/api/auth/me":
            user_id = qs.get("user_id", [""])[0]
            email = qs.get("email", [""])[0]
            if not user_id:
                # Try token as fallback
                token = extract_token_from_headers(self.headers)
                if token:
                    try:
                        import jwt as pyjwt
                        unverified = pyjwt.decode(token, options={"verify_signature": False})
                        user_id = unverified.get("sub", "")
                        email = unverified.get("email", "") or email
                    except: pass
            if not user_id:
                self.send_json({"error": "No user_id"}, 401); return
            try:
                creator = get_creator_by_clerk_id(user_id)
                if not creator and email:
                    # Fallback: find by email (for creators made before auth system)
                    creator = get_creator_by_email(email)
                    if creator:
                        # Link clerk_user_id to existing creator
                        save_creator(creator["slug"], {"clerk_user_id": user_id})
                        creator = get_creator(creator["slug"])
                if not creator and email:
                    slug = email.split("@")[0].lower().replace(".", "-").replace("_", "-")
                    base_slug = slug
                    counter = 1
                    while get_creator(slug):
                        slug = f"{base_slug}-{counter}"
                        counter += 1
                    save_creator(slug, {"clerk_user_id": user_id, "email": email, "channel_name": slug})
                    creator = get_creator(slug)
                self.send_json({"user": {"user_id": user_id, "email": email}, "creator": creator or {}})
            except Exception as e:
                import traceback
                print(f"❌ auth/me error: {traceback.format_exc()}")
                self.send_json({"error": str(e)}, 500)

    def do_POST(self):
        p = urlparse(self.path)

        # ── Raw body endpoints (must be read before read_body()) ──────────────
        if p.path == "/api/stripe/webhook":
            length = int(self.headers.get("Content-Length", 0))
            raw_body = self.rfile.read(length)
            sig_header = self.headers.get("Stripe-Signature", "")
            try:
                event = handle_stripe_webhook(raw_body, sig_header)
                event_type = event.get("type", "")
                obj = event.get("data", {}).get("object", {})
                if event_type == "checkout.session.completed":
                    meta = obj.get("metadata", {})
                    slug = meta.get("slug", "")
                    plan = meta.get("plan", "creator")
                    if slug:
                        update_subscription(
                            slug=slug,
                            stripe_customer_id=obj.get("customer", ""),
                            stripe_subscription_id=obj.get("subscription", ""),
                            status="active", plan=plan
                        )
                        print(f"✅ Subscription aktiviert: {slug} → {plan}")
                elif event_type in ("customer.subscription.deleted",):
                    customer_id = obj.get("customer", "")
                    if customer_id:
                        update_subscription_by_customer(customer_id, "inactive", "free")
                self.send_json({"ok": True})
            except Exception as e:
                print(f"⚠️ Webhook error: {e}")
                self.send_json({"error": str(e)}, 400)
            return

        if p.path == "/api/clerk/webhook":
            length = int(self.headers.get("Content-Length", 0))
            raw_body = self.rfile.read(length)
            try:
                payload = json.loads(raw_body)
                result = handle_clerk_webhook(payload)
                self.send_json(result)
            except Exception as e:
                self.send_json({"error": str(e)}, 400)
            return

        data = self.read_body()

        if p.path == "/api/creator/save":
            slug = data.get("slug", "").lower().replace(" ", "-")
            if not slug:
                self.send_json({"error": "No slug"}, 400)
                return
            save_creator(slug, data)
            self.send_json({"ok": True, "slug": slug, "url": f"/c/{slug}"})

        elif p.path == "/api/chat":
            slug = data.get("slug", "")
            messages = data.get("messages", [])
            session_id = data.get("session_id", "anon")

            if not slug or not messages:
                self.send_json({"error": "Missing data"}, 400)
                return

            # Load session from DB
            session_data = get_session(session_id)
            if session_data and session_data.get("diag_state"):
                ds = session_data["diag_state"]
                if isinstance(ds, str):
                    ds = json.loads(ds)
                session_data["diag_state"] = ds

            try:
                result = get_ai_response(slug, messages, session_data or {})
                save_session(
                    session_id=session_id,
                    slug=slug,
                    message_count=len(messages),
                    is_warm=bool(result.get("warm_lead")),
                    diag_state=result.get("diag_state", {})
                )
                self.send_json(result)
            except Exception as e:
                self.send_json({"response": str(e), "error": str(e)})

        elif p.path == "/api/quiz":
            try:
                slug = data.get("slug", "")
                difficulty = int(data.get("difficulty", 1))
                knowledge = load_knowledge(slug)
                # Fallback: use config bio + product descriptions if no transcript
                if not knowledge:
                    config = load_json(creator_dir(slug) / "config.json", {})
                    parts = []
                    if config.get("bio"): parts.append("Creator bio: " + config["bio"])
                    for p2 in config.get("products", []):
                        if p2.get("name"): parts.append("Product: " + p2["name"])
                    knowledge = "\n".join(parts)
                if not knowledge:
                    self.send_json({"error": "No knowledge base loaded yet. Please add content in the creator setup first."}, 400)
                    return
                diff_map = {
                    1: ("LEICHT", "Ask about a concrete, practical concept a beginner finds useful. Wrong answers clearly different."),
                    2: ("MITTEL", "Ask about HOW or WHY something works. All wrong answers sound plausible to someone who half-understood."),
                    3: ("SCHWER", "Ask about a subtle distinction or common expert mistake. All 4 options look correct to surface-level knowers.")
                }
                diff_label, diff_instr = diff_map.get(difficulty, diff_map[1])
                api_key_val = os.environ.get("ANTHROPIC_API_KEY","")
                if not api_key_val:
                    self.send_json({"error": "API Key nicht gesetzt. Server neu starten mit: export ANTHROPIC_API_KEY=sk-ant-..."}, 500)
                    return
                client = anthropic.Anthropic(api_key=api_key_val)
                quiz_resp = client.messages.create(
                    model="claude-sonnet-4-20250514",
                    max_tokens=500,
                    system=(
                        "You are a quiz generator for an interactive learning app. "
                        "Generate ONE multiple choice question testing the learner on the creator's content. "
                        "RULES: Ask about a real concept/technique/insight from the knowledge base. "
                        "NEVER ask meta-questions about the channel name or creator identity. "
                        "Questions must teach something practical. "
                        "Generate in the SAME LANGUAGE as the knowledge base (German if German). "
                        f"Difficulty {difficulty}/3 ({diff_label}): {diff_instr} "
                        "Return ONLY valid JSON with NO markdown fences: "
                        '{"question":"...","options":["A) ...","B) ...","C) ...","D) ..."],"correct":"A","explanation":"why correct in 1-2 sentences"} '
                        f"Knowledge base:\n{knowledge[:8000]}"
                    ),
                    messages=[{"role":"user","content":f"Generate a {diff_label} quiz question now."}]
                )
                raw = quiz_resp.content[0].text.strip()
                raw = re.sub(r"```json|```","",raw).strip()
                quiz = json.loads(raw)
                self.send_json({"quiz": quiz})
            except json.JSONDecodeError as e:
                self.send_json({"error": f"JSON parse error: {e}", "raw": raw[:300]}, 500)
            except Exception as e:
                import traceback
                self.send_json({"error": str(e), "trace": traceback.format_exc()[-500:]}, 500)

        elif p.path == "/api/quiz-explain":
            slug = data.get("slug", "")
            question = data.get("question", "")
            correct_answer = data.get("correct_answer", "")
            knowledge = load_knowledge(slug)
            try:
                client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY",""))
                exp_resp = client.messages.create(
                    model="claude-sonnet-4-20250514",
                    max_tokens=200,
                    system=f"You explain quiz answers briefly in 1-2 sentences. Be clear, educational and encouraging. Knowledge base: {knowledge[:3000]}",
                    messages=[{"role":"user","content":f"Question: {question}\nCorrect answer: {correct_answer}\nExplain why this is correct in 1-2 sentences."}]
                )
                self.send_json({"explanation": exp_resp.content[0].text.strip()})
            except Exception as e:
                self.send_json({"explanation": ""})

        elif p.path == "/api/knowledge/append":
            slug = data.get("slug", "")
            new_text = data.get("text", "").strip()
            source = data.get("source", "manual")
            if not slug or not new_text:
                self.send_json({"error": "Missing slug or text"}, 400)
                return
            transcript_path = creator_dir(slug) / "transcript.txt"
            try:
                timestamp = datetime.now().strftime("%d.%m.%Y %H:%M")
                separator = f"\n\n{'='*60}\n## MANUELL HINZUGEFÜGT ({source.upper()}) — {timestamp}\n{'='*60}\n\n"
                with open(transcript_path, "a", encoding="utf-8") as f:
                    f.write(separator + new_text + "\n")
                # Update word count in config
                cfg = load_json(creator_dir(slug) / "config.json", {})
                total_text = transcript_path.read_text(encoding="utf-8")
                cfg["word_count"] = len(total_text.split())
                cfg["last_knowledge_update"] = timestamp
                save_json(creator_dir(slug) / "config.json", cfg)
                self.send_json({"ok": True, "words_added": len(new_text.split())})
            except Exception as e:
                self.send_json({"error": str(e)}, 500)

        elif p.path == "/api/knowledge/upload":
            ctype = self.headers.get("Content-Type", "")
            if "multipart/form-data" not in ctype:
                self.send_json({"error": "Expected multipart/form-data"}, 400)
                return
            try:
                import io as _io
                clen = int(self.headers.get("Content-Length", 0))
                fields = parse_multipart(self.rfile, ctype, clen)
                slug = fields.get("slug", "")
                source = fields.get("source", "file")
                file_info = fields.get("file", {})
                if not slug or not isinstance(file_info, dict) or not file_info.get("filename"):
                    self.send_json({"error": "Missing slug or file"}, 400)
                    return
                raw = file_info["data"]
                fname = file_info["filename"]
                fname_lower = fname.lower()
                text = ""
                method = "unknown"

                if fname_lower.endswith((".txt", ".md", ".csv", ".rtf")):
                    try: text = raw.decode("utf-8")
                    except: text = raw.decode("latin-1", errors="ignore")
                    method = "plaintext"

                elif fname_lower.endswith(".pdf"):
                    try:
                        import pypdf
                        reader = pypdf.PdfReader(_io.BytesIO(raw))
                        pages = [pg.extract_text() or "" for pg in reader.pages]
                        text = "\n\n".join(pg.strip() for pg in pages if pg.strip())
                        method = f"pypdf ({len(reader.pages)} pages)"
                    except Exception as pdf_err:
                        import re as _re
                        decoded = raw.decode("latin-1", errors="ignore")
                        runs = _re.findall(r'[A-Za-zÄÖÜäöüß][A-Za-zÄÖÜäöüß0-9\s,.\-:!?()]{12,}', decoded)
                        text = " ".join(runs)
                        method = f"pdf-regex ({pdf_err})"

                elif fname_lower.endswith(".docx"):
                    try:
                        import docx as _docx
                        doc = _docx.Document(_io.BytesIO(raw))
                        paras = [p.text for p in doc.paragraphs if p.text.strip()]
                        for tbl in doc.tables:
                            for row in tbl.rows:
                                for cell in row.cells:
                                    if cell.text.strip(): paras.append(cell.text.strip())
                        text = "\n".join(paras)
                        method = "python-docx"
                    except Exception as docx_err:
                        import zipfile, re as _re
                        try:
                            z = zipfile.ZipFile(_io.BytesIO(raw))
                            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
                            text = " ".join(_re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml))
                            method = "docx-xml-fallback"
                        except:
                            text = raw.decode("utf-8", errors="ignore")
                            method = "docx-raw"

                elif fname_lower.endswith((".html", ".htm")):
                    import re as _re
                    decoded = raw.decode("utf-8", errors="ignore")
                    text = _re.sub(r'<[^>]+>', ' ', decoded)
                    text = _re.sub(r'\s+', ' ', text).strip()
                    method = "html-strip"

                elif fname_lower.endswith(".json"):
                    import json as _json
                    try: text = _json.dumps(_json.loads(raw.decode("utf-8")), ensure_ascii=False, indent=2)
                    except: text = raw.decode("utf-8", errors="ignore")
                    method = "json"

                else:
                    text = raw.decode("utf-8", errors="ignore")
                    method = "raw-utf8"

                text = text.strip()
                if not text or len(text) < 30:
                    self.send_json({"error": f"Kein Text aus '{fname}' extrahierbar (Methode: {method}). Bitte als TXT exportieren."}, 400)
                    return

                if len(text) > 200000:
                    text = text[:200000] + "\n[... bei 200k Zeichen abgeschnitten]"

                transcript_path = creator_dir(slug) / "transcript.txt"
                timestamp = datetime.now().strftime("%d.%m.%Y %H:%M")
                sep = f"\n\n{'='*60}\n## {fname} ({source.upper()}) — {timestamp}\n{'='*60}\n\n"
                with open(transcript_path, "a", encoding="utf-8") as f2:
                    f2.write(sep + text + "\n")
                cfg2 = load_json(creator_dir(slug) / "config.json", {})
                cfg2["last_knowledge_update"] = timestamp
                cfg2["has_knowledge"] = True
                save_json(creator_dir(slug) / "config.json", cfg2)
                self.send_json({"ok": True, "words": len(text.split()), "chars": len(text), "filename": fname, "method": method})
            except Exception as e:
                import traceback
                self.send_json({"error": str(e), "trace": traceback.format_exc()[-500:]}, 500)

        elif p.path == "/api/greeting-video/upload":
            # Handle multipart form upload
            import cgi
            ctype = self.headers.get("Content-Type", "")
            if "multipart/form-data" not in ctype:
                self.send_json({"error": "Expected multipart"}, 400)
                return
            try:
                form = cgi.FieldStorage(
                    fp=self.rfile,
                    headers=self.headers,
                    environ={"REQUEST_METHOD": "POST", "CONTENT_TYPE": ctype}
                )
                slug = form.getvalue("slug", "")
                video_field = form["video"]
                if not slug or not video_field.filename:
                    self.send_json({"error": "Missing slug or video"}, 400)
                    return
                ext = Path(video_field.filename).suffix.lower() or ".mp4"
                video_path = creator_dir(slug) / f"greeting{ext}"
                with open(video_path, "wb") as f:
                    f.write(video_field.file.read())
                # Save reference in config
                cfg = load_json(creator_dir(slug) / "config.json", {})
                cfg["greeting_video"] = f"greeting{ext}"
                save_json(creator_dir(slug) / "config.json", cfg)
                self.send_json({"ok": True, "file": f"greeting{ext}"})
            except Exception as e:
                self.send_json({"error": str(e)}, 500)

        elif p.path == "/api/leads":
            slug = data.get("slug", "")
            leads = get_leads(slug)
            leads_dummy = leads  # alias
            self.send_json({"leads": leads, "total": len(sessions), "warm": len(leads)})

        elif p.path == "/api/stripe/checkout":
            slug = data.get("slug", "")
            plan = data.get("plan", "creator")
            clerk_user_id = data.get("clerk_user_id", "")
            email = data.get("email", "")
            base_url = data.get("base_url", "http://localhost:7891")
            success_slug = data.get("success_slug", slug)
            try:
                url = create_checkout_session(
                    slug=slug, plan=plan,
                    clerk_user_id=clerk_user_id, email=email,
                    success_url=f"{base_url}/admin?slug={success_slug}&payment=success",
                    cancel_url=f"{base_url}/admin?slug={success_slug}&payment=cancelled"
                )
                self.send_json({"ok": True, "url": url})
            except Exception as e:
                self.send_json({"error": str(e)}, 400)

        elif p.path == "/api/stripe/portal":
            slug = data.get("slug", "")
            base_url = data.get("base_url", "http://localhost:7891")
            creator = get_creator(slug)
            if not creator or not creator.get("stripe_customer_id"):
                self.send_json({"error": "No Stripe customer found"}, 400)
                return
            try:
                url = create_billing_portal_session(
                    creator["stripe_customer_id"],
                    return_url=f"{base_url}/admin?slug={slug}"
                )
                self.send_json({"ok": True, "url": url})
            except Exception as e:
                self.send_json({"error": str(e)}, 400)



        else:
            self.send_response(404)
            self.end_headers()


# ── CREATOR SETUP HTML ────────────────────────────────────────────────────────
LOGIN_HTML = """<!DOCTYPE html>
<html lang="de">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>KnowFlow — Login</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&family=DM+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
:root{--bg:#000;--s:#111;--s2:#1a1a1a;--b:rgba(255,255,255,.08);--b2:rgba(255,255,255,.13);--acc:#ff6830;--acc2:#ff9f0a;--grad:linear-gradient(135deg,#ff6830,#ff9f0a);--t:#f0f0f0;--t2:#999;--t3:#555}
html,body{background:var(--bg);color:var(--t);font-family:'DM Sans',sans-serif;min-height:100vh;display:flex;flex-direction:column;align-items:center;justify-content:center;-webkit-font-smoothing:antialiased}
body{background-image:radial-gradient(ellipse 70% 40% at 50% 0%,rgba(255,104,48,.07) 0%,transparent 60%)}
.card{background:var(--s);border:1px solid var(--b);border-radius:20px;padding:40px 36px;width:100%;max-width:420px;text-align:center}
.logo{font-family:'DM Mono',monospace;font-size:20px;font-weight:500;margin-bottom:28px}
.logo em{background:var(--grad);-webkit-background-clip:text;-webkit-text-fill-color:transparent;font-style:normal}
h2{font-size:22px;font-weight:700;margin-bottom:8px}
.sub{font-size:13px;color:var(--t3);margin-bottom:32px;line-height:1.5}
#clerk-mount{width:100%}
.cl-rootBox{width:100%}
.cl-card{background:transparent!important;box-shadow:none!important;border:none!important;padding:0!important}
.cl-headerTitle,.cl-headerSubtitle{display:none!important}
.cl-socialButtonsBlockButton{background:var(--s2)!important;border:1px solid var(--b2)!important;color:var(--t)!important;border-radius:10px!important}
.cl-formButtonPrimary{background:linear-gradient(135deg,#ff6830,#ff9f0a)!important;border-radius:10px!important;font-weight:700!important}
.cl-formFieldInput{background:var(--s2)!important;border:1px solid var(--b2)!important;color:var(--t)!important;border-radius:8px!important}
.cl-footerActionLink{color:var(--acc)!important}
</style>
</head>
<body>
<div class="card">
  <div class="logo">Know<em>Flow</em></div>
  <h2>Creator Login</h2>
  <p class="sub">Melde dich an um dein KnowFlow Dashboard zu verwalten.</p>
  <div id="clerk-mount"></div>
</div>
<script>window.__clerk_pk="{{CLERK_PK}}";</script>
<script async crossorigin="anonymous"
  data-clerk-publishable-key="{{CLERK_PK}}"
  src="https://cdn.jsdelivr.net/npm/@clerk/clerk-js@5/dist/clerk.browser.js">
</script>
<script>
window.addEventListener('load', async () => {
  const pk = window.__clerk_pk;
  if (!pk || !pk.startsWith('pk_')) {
    document.getElementById('clerk-mount').innerHTML =
      '<p style="color:#ff453a;font-size:13px;margin-top:8px">CLERK_PUBLISHABLE_KEY fehlt in .env</p>';
    return;
  }
  await window.Clerk.load();
  if (window.Clerk.user) {
    redirectToAdmin();
    return;
  }
  window.Clerk.mountSignIn(document.getElementById('clerk-mount'), {
    afterSignInUrl: '/auth/callback',
    afterSignUpUrl: '/auth/callback',
  });
});
async function redirectToAdmin() {
  try {
    const token = await window.Clerk.session.getToken();
    const res = await fetch('/api/auth/me', { headers: {'Authorization':'Bearer '+token} });
    const d = await res.json();
    if (d.creator && d.creator.slug) { window.location.href = '/admin?slug=' + d.creator.slug; return; }
  } catch(e) {}
  const email = window.Clerk.user?.primaryEmailAddress?.emailAddress || '';
  window.location.href = '/admin?slug=' + email.split('@')[0].toLowerCase().replace(/[^a-z0-9]/g,'-');
}
</script>
</body>
</html>"""

CALLBACK_HTML = """<!DOCTYPE html>
<html><head><meta charset="UTF-8"><title>KnowFlow</title>
<meta name="clerk-pk" content="{{CLERK_PK}}">
<style>*{box-sizing:border-box;margin:0;padding:0}body{background:#000;color:#f0f0f0;font-family:'DM Sans',sans-serif;display:flex;align-items:center;justify-content:center;min-height:100vh;flex-direction:column;gap:16px}.spin{width:32px;height:32px;border:3px solid rgba(255,255,255,.1);border-top-color:#ff6830;border-radius:50%;animation:spin .7s linear infinite}@keyframes spin{to{transform:rotate(360deg)}}p{font-size:14px;color:#555}</style>
</head><body>
<div class="spin"></div><p>Einen Moment...</p>
<script async crossorigin="anonymous"
  data-clerk-publishable-key="{{CLERK_PK}}"
  src="https://cdn.jsdelivr.net/npm/@clerk/clerk-js@5/dist/clerk.browser.js">
</script>
<script>
window.addEventListener('load', async () => {
  const pk = document.querySelector('meta[name=clerk-pk]')?.content||'';
  if (!pk) { window.location.href='/login'; return; }
  await window.Clerk.load();
  const user = window.Clerk.user;
  if (!user) { window.location.href='/login'; return; }
  const email = user.primaryEmailAddress?.emailAddress||'';
  const userId = user.id||'';
  const slug = email.split('@')[0].toLowerCase().replace(/[^a-z0-9]/g,'-');
  try {
    const res = await fetch(`/api/auth/me?user_id=${encodeURIComponent(userId)}&email=${encodeURIComponent(email)}`);
    const d = await res.json();
    if (d.creator && d.creator.slug) { window.location.href='/admin?slug='+d.creator.slug; return; }
  } catch(e) {}
  window.location.href='/admin?slug='+slug;
});
</script>
</body></html>"""


CREATOR_HTML = r"""<!DOCTYPE html>
<html lang="de">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>KnowFlow — Creator Setup</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:ital,opsz,wght@0,9..40,300;0,9..40,400;0,9..40,500;0,9..40,600;0,9..40,700&family=DM+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
:root{
  --bg:#000;--s:#111;--s2:#1a1a1a;--s3:#222;--b:rgba(255,255,255,.08);--b2:rgba(255,255,255,.13);
  --acc:#ff6830;--acc-bg:rgba(255,104,48,.09);--acc-b:rgba(255,104,48,.3);
  --acc2:#ff9f0a;
  --blue:#4f8eff;--blue-bg:rgba(79,142,255,.08);--blue-b:rgba(79,142,255,.2);
  --t:#f0f0f0;--t2:#999;--t3:#555;
  --ok:#30d158;--err:#ff453a;--warn:#ff9f0a;
  --grad:linear-gradient(135deg,#ff6830,#ff9f0a);
}
html{background:var(--bg)}
body{
  background:var(--bg);color:var(--t);font-family:'DM Sans',sans-serif;min-height:100vh;-webkit-font-smoothing:antialiased;
  background-image:radial-gradient(ellipse 70% 50% at 50% 100%,rgba(255,104,48,.12) 0%,transparent 60%),
    radial-gradient(ellipse 40% 30% at 80% 80%,rgba(255,159,10,.07) 0%,transparent 50%);
}
nav{
  position:sticky;top:0;z-index:50;
  background:rgba(0,0,0,.9);backdrop-filter:blur(20px);
  border-bottom:1px solid var(--b);
  padding:14px 40px;display:flex;align-items:center;justify-content:space-between;
}
.logo{font-family:'DM Mono',monospace;font-size:16px;font-weight:500;letter-spacing:.05em}
.logo span{color:var(--t2)}
.nav-right{display:flex;align-items:center;gap:16px}
.nav-link{color:var(--t3);font-size:13px;text-decoration:none;transition:color .2s}
.nav-link:hover{color:var(--t)}
.nav-btn{
  background:var(--s2);border:1px solid var(--b2);
  color:var(--t2);font-size:12px;font-weight:600;
  padding:7px 14px;cursor:pointer;font-family:'DM Sans',sans-serif;
  border-radius:8px;transition:all .15s;text-decoration:none;
}
.nav-btn:hover{border-color:var(--acc-b);color:var(--acc)}
.container{max-width:720px;margin:0 auto;padding:48px 24px 120px}
.hero-badge{
  display:inline-flex;align-items:center;gap:8px;
  font-family:'DM Mono',monospace;font-size:10px;letter-spacing:.1em;text-transform:uppercase;
  color:var(--acc);border:1px solid var(--acc-b);
  padding:4px 12px;border-radius:20px;margin-bottom:20px;
}
h1{font-size:clamp(36px,6vw,56px);font-weight:700;letter-spacing:-.03em;line-height:.95;margin-bottom:14px}
h1 em{background:var(--grad);-webkit-background-clip:text;-webkit-text-fill-color:transparent;font-style:normal}
.sub{color:var(--t2);font-size:15px;max-width:480px;line-height:1.7;margin-bottom:48px}

/* STEPS */
.steps{display:flex;gap:0;margin-bottom:40px;position:relative}
.steps::before{content:'';position:absolute;top:19px;left:20px;right:20px;height:1px;background:var(--b)}
.step{flex:1;text-align:center;position:relative;z-index:1}
.step-dot{
  width:38px;height:38px;border-radius:50%;background:var(--bg);
  border:1px solid var(--b);margin:0 auto 10px;
  display:flex;align-items:center;justify-content:center;
  font-family:'DM Mono',monospace;font-size:12px;color:var(--t3);
  transition:all .25s;
}
.step.active .step-dot{border-color:var(--acc);color:var(--acc);background:var(--acc-bg);box-shadow:0 0 16px rgba(255,104,48,.18)}
.step.done .step-dot{border-color:var(--ok);color:var(--ok);background:rgba(48,209,88,.08)}
.step-label{font-size:10px;color:var(--t3);font-family:'DM Mono',monospace;letter-spacing:.04em}
.step.active .step-label{color:var(--acc)}
.step.done .step-label{color:var(--ok)}

/* CARDS */
.card{background:var(--s);border:1px solid var(--b);border-radius:16px;padding:32px;margin-bottom:16px;display:none}
.card.active{display:block}
.card-title{font-size:20px;font-weight:700;letter-spacing:-.02em;margin-bottom:6px}
.card-sub{color:var(--t2);font-size:13px;margin-bottom:28px;line-height:1.65}

label{display:block;font-family:'DM Mono',monospace;font-size:10px;letter-spacing:.1em;
  text-transform:uppercase;color:var(--t3);margin-bottom:6px}
input,textarea,select{
  width:100%;background:var(--bg);border:1px solid var(--b);
  color:var(--t);font-family:'DM Sans',sans-serif;font-size:14px;
  padding:11px 14px;outline:none;transition:border-color .2s;border-radius:10px;resize:vertical;
}
input:focus,textarea:focus,select:focus{border-color:var(--acc-b)}
input::placeholder,textarea::placeholder{color:var(--t3)}
.field{margin-bottom:18px}
.row2{display:grid;grid-template-columns:1fr 1fr;gap:14px;margin-bottom:18px}

/* PRODUCTS */
.product-list{display:flex;flex-direction:column;gap:9px;margin-bottom:14px}
.product-item{
  background:var(--s2);border:1px solid var(--b);border-radius:10px;
  padding:14px;display:grid;grid-template-columns:1fr 1fr auto;gap:10px;align-items:end;
}
.product-item input{background:var(--bg)}
.btn-remove{background:none;border:none;color:var(--err);cursor:pointer;font-size:18px;padding:6px;transition:opacity .2s;width:32px;text-align:center}
.btn-remove:hover{opacity:.6}
.btn-add{
  background:none;border:1px dashed var(--b);color:var(--t3);width:100%;
  padding:11px;cursor:pointer;font-family:'DM Sans',sans-serif;font-size:13px;
  transition:all .2s;border-radius:10px;
}
.btn-add:hover{border-color:var(--acc-b);color:var(--acc)}

/* BUTTONS */
.btn{
  padding:12px 28px;font-family:'DM Sans',sans-serif;font-weight:700;font-size:14px;
  cursor:pointer;border:none;transition:all .18s;border-radius:10px;letter-spacing:.01em;
}
.btn-primary{background:var(--grad);color:#fff;border-color:transparent!important}
.btn-primary:hover{background:#d4f850;transform:translateY(-1px)}
.btn-ghost{background:none;border:1px solid var(--b);color:var(--t2)}
.btn-ghost:hover{border-color:var(--b2);color:var(--t)}
.btn:disabled{opacity:.35;cursor:not-allowed;transform:none!important}
.btn-row{display:flex;gap:10px;margin-top:28px;flex-wrap:wrap}

/* LOG */
.log-card{background:var(--bg);border:1px solid var(--b);border-radius:12px;margin-top:18px;display:none;overflow:hidden}
.log-header{
  padding:10px 14px;border-bottom:1px solid var(--b);
  display:flex;align-items:center;gap:8px;
  font-family:'DM Mono',monospace;font-size:10px;color:var(--t3);
}
.log-dot{width:6px;height:6px;border-radius:50%;background:var(--acc);animation:pulse 1s ease-in-out infinite;box-shadow:0 0 6px var(--acc)}
.log-dot.done{animation:none;background:var(--ok)}
@keyframes pulse{0%,100%{opacity:1}50%{opacity:.2}}
#log-output{padding:14px;font-family:'DM Mono',monospace;font-size:11px;line-height:1.8;max-height:280px;overflow-y:auto}
.l-info{color:var(--t)}.l-success{color:var(--ok)}.l-warn{color:var(--warn)}
.l-error{color:var(--err)}.l-dim{color:var(--t3)}.l-done{color:var(--acc);font-weight:700}
.l-file{color:var(--blue)}.l-start{color:var(--acc);font-weight:600}
.pbar{height:2px;background:var(--b);display:none;border-radius:1px;margin-top:2px}
.pfill{height:100%;background:var(--grad);width:0;transition:width .4s;border-radius:1px}

/* KNOWLEDGE SOURCES */
.sources-list{display:flex;flex-direction:column;gap:9px;margin-bottom:14px}
.source-item{
  background:var(--s2);border:1px solid var(--b);border-radius:10px;
  padding:14px;position:relative;
}
.source-item.added{border-color:rgba(48,209,88,.3);background:rgba(48,209,88,.04)}
.source-header{display:flex;align-items:center;justify-content:space-between;margin-bottom:10px}
.source-type{
  font-family:'DM Mono',monospace;font-size:10px;font-weight:500;
  letter-spacing:.06em;text-transform:uppercase;
  color:var(--acc);
}
.source-badge{
  font-family:'DM Mono',monospace;font-size:9px;font-weight:500;
  padding:2px 8px;border-radius:20px;
  background:rgba(48,209,88,.1);border:1px solid rgba(48,209,88,.2);color:var(--ok);
  display:none;
}
.source-item.added .source-badge{display:inline}
.source-textarea{min-height:90px}
.char-count{font-family:'DM Mono',monospace;font-size:10px;color:var(--t3);text-align:right;margin-top:4px}

/* FILE UPLOAD */
.upload-area{
  border:1px dashed var(--b2);border-radius:10px;padding:24px;
  text-align:center;cursor:pointer;transition:all .2s;
  position:relative;overflow:hidden;
}
.upload-area:hover,.upload-area.drag{border-color:var(--acc-b);background:var(--acc-bg)}
.upload-area input[type=file]{
  position:absolute;inset:0;opacity:0;cursor:pointer;
  width:100%;height:100%;
}
.upload-icon{font-size:28px;margin-bottom:8px}
.upload-text{font-size:13px;color:var(--t2);margin-bottom:4px}
.upload-sub{font-size:11px;color:var(--t3);font-family:'DM Mono',monospace}
.upload-preview{margin-top:10px;display:none}
.upload-preview video{max-width:100%;border-radius:8px;max-height:180px}
.upload-preview-name{font-size:12px;color:var(--ok);font-family:'DM Mono',monospace;margin-top:6px}

/* FILE DROP ZONES */
.file-drop-row{display:grid;grid-template-columns:1fr 1fr;gap:8px;margin-bottom:4px}
.file-drop-zone{
  border:1px dashed var(--b2);border-radius:10px;padding:18px 10px;
  text-align:center;cursor:pointer;transition:all .2s;
  background:var(--bg);
}
.file-drop-zone:hover{border-color:var(--acc-b);background:var(--acc-bg)}
.file-drop-zone.done{border-color:var(--ok);border-style:solid;background:rgba(48,209,88,.04)}
.fdz-icon{font-size:22px;margin-bottom:6px}
.fdz-label{font-size:12px;font-weight:600;color:var(--t1);margin-bottom:2px}
.fdz-sub{font-size:10px;color:var(--t3);font-family:'DM Mono',monospace}
.fdz-file{font-size:10px;color:var(--ok);margin-top:4px;font-family:'DM Mono',monospace;word-break:break-all}

/* UPLOADED FILES LIST */
#knowledge-files-list{display:flex;flex-direction:column;gap:6px;margin-bottom:12px}
.kf-item{
  display:flex;align-items:center;gap:10px;
  background:rgba(48,209,88,.05);border:1px solid rgba(48,209,88,.2);
  border-radius:8px;padding:8px 12px;font-size:12px;
}
.kf-icon{font-size:16px;flex-shrink:0}
.kf-name{flex:1;color:var(--t1);font-weight:500}
.kf-words{font-size:10px;color:var(--ok);font-family:'DM Mono',monospace;flex-shrink:0}

/* KB STATUS */
.kb-status{
  background:var(--bg);border:1px solid var(--b);border-radius:10px;padding:14px;
  display:flex;align-items:center;gap:12px;
}
.kb-icon{font-size:22px}
.kb-info{flex:1}
.kb-title{font-size:13px;font-weight:600}
.kb-sub{font-size:11px;color:var(--t3);font-family:'DM Mono',monospace;margin-top:2px}
.kb-dot{width:8px;height:8px;border-radius:50%;background:var(--ok);flex-shrink:0}
.kb-dot.pending{background:var(--warn);animation:pulse 1.5s ease-in-out infinite}

/* SUCCESS */
.success-box{background:rgba(48,209,88,.04);border:1px solid rgba(48,209,88,.2);border-radius:12px;padding:22px;margin-top:18px;display:none}
.success-box h3{color:var(--ok);font-size:17px;margin-bottom:6px}
.success-url{
  font-family:'DM Mono',monospace;font-size:12px;color:var(--blue);
  background:var(--bg);padding:10px 13px;margin-top:10px;border-radius:8px;
  border:1px solid var(--b);cursor:pointer;word-break:break-all;
}
.success-url:hover{background:var(--s2)}
.slug-preview{font-family:'DM Mono',monospace;font-size:11px;color:var(--t3);margin-top:5px}
.slug-preview span{color:var(--acc)}

/* DASHBOARD (step 5) */
.dash-grid{display:grid;grid-template-columns:1fr 1fr;gap:12px;margin-bottom:20px}
.dash-card{
  background:var(--s2);border:1px solid var(--b);border-radius:12px;
  padding:18px;cursor:pointer;transition:all .18s;text-decoration:none;color:var(--t);
}
.dash-card:hover{border-color:var(--b2)}
.dash-card-icon{font-size:24px;margin-bottom:10px}
.dash-card-title{font-size:14px;font-weight:700;margin-bottom:4px}
.dash-card-sub{font-size:12px;color:var(--t3);line-height:1.5}
.sep{height:1px;background:var(--b);margin:24px 0}

@media(max-width:600px){
  .row2{grid-template-columns:1fr}
  .product-item{grid-template-columns:1fr auto}
  .dash-grid{grid-template-columns:1fr}
  nav{padding:14px 20px}
  .container{padding:32px 16px 80px}
}
</style>
</head>
<body>

<nav>
  <div class="logo">Know<span>Flow</span></div>
  <div class="nav-right">
    <a class="nav-btn" href="#" id="dashboard-link" style="display:none" onclick="showDashboard()">📊 Mein Dashboard</a>
    <a class="nav-link" href="#" onclick="showStep(1)">← Setup</a>
  </div>
</nav>

<div class="container">
  <div class="hero-badge">Creator Setup</div>
  <h1>Deine <em>KnowFlow</em><br>Seite einrichten.</h1>
  <p class="sub">YouTube-Channel verbinden, Wissensquellen hinzufügen, Produkte konfigurieren — dein AI Sales Funnel ist in Minuten live.</p>

  <!-- STEPS -->
  <div class="steps">
    <div class="step active" id="step-1"><div class="step-dot">1</div><div class="step-label">Channel</div></div>
    <div class="step" id="step-2"><div class="step-dot">2</div><div class="step-label">Profil</div></div>
    <div class="step" id="step-3"><div class="step-dot">3</div><div class="step-label">Produkte</div></div>
    <div class="step" id="step-4"><div class="step-dot">4</div><div class="step-label">Live</div></div>
  </div>

  <!-- STEP 1: CHANNEL + VIDEO -->
  <div class="card active" id="card-1">
    <div class="card-title">🎬 YouTube Channel verbinden</div>
    <div class="card-sub">Wir scannen deinen Channel und transkribieren alle Videos in eine durchsuchbare Wissensbasis — das Gehirn deines AI-Assistenten.</div>

    <div class="field">
      <label>YouTube Channel URL</label>
      <input type="text" id="channel-url" placeholder="https://www.youtube.com/@DeinChannel">
    </div>
    <div class="field">
      <label>Deine Seiten-URL (Slug)</label>
      <input type="text" id="creator-slug" placeholder="z.B. svengold" oninput="updateSlugPreview()">
      <div class="slug-preview">Deine Seite: knowflow.io/c/<span id="slug-display">deinname</span></div>
    </div>
    <div class="row2">
      <div class="field">
        <label>Max. Videos transkribieren</label>
        <select id="max-videos">
          <option value="10">10 Videos (schneller Test)</option>
          <option value="30" selected>30 Videos (empfohlen)</option>
          <option value="50">50 Videos</option>
          <option value="100">100 Videos</option>
          <option value="0">Alle Videos</option>
        </select>
      </div>
    </div>

    <div class="sep" style="margin:20px 0 20px"></div>
    <div class="card-title" style="font-size:16px;margin-bottom:6px">🎥 Begrüßungsvideo <span style="font-size:12px;font-weight:400;color:var(--t3)">(optional)</span></div>
    <div class="card-sub" style="margin-bottom:16px">Füge einen YouTube- oder Vimeo-Link ein — auch nicht-gelistete Videos funktionieren. Wird ganz oben auf deiner Seite als persönliche Begrüßung angezeigt.</div>
    <div class="field" style="margin-bottom:0">
      <div style="display:flex;gap:8px;align-items:center">
        <input type="text" id="greeting-video-url" placeholder="https://youtube.com/watch?v=... oder https://vimeo.com/..."
          style="flex:1" oninput="previewVideoUrl(this.value)">
        <button class="btn btn-ghost" style="flex-shrink:0;font-size:12px;padding:8px 14px" onclick="previewVideoUrl(document.getElementById('greeting-video-url').value)">Vorschau</button>
      </div>
    </div>
    <div id="video-preview" style="display:none;margin-top:12px">
      <div style="position:relative;width:100%;aspect-ratio:16/9;border-radius:10px;overflow:hidden;background:#000">
        <iframe id="video-iframe" src="" frameborder="0" allowfullscreen
          style="position:absolute;inset:0;width:100%;height:100%"></iframe>
      </div>
      <div id="video-name" style="font-size:11px;color:var(--ok);margin-top:6px;font-family:'DM Mono',monospace"></div>
    </div>

    <div class="sep" style="margin:24px 0 20px"></div>
    <div class="card-title" style="font-size:16px;margin-bottom:4px">🧠 Wissensquellen <span style="font-size:12px;font-weight:400;color:var(--t3)">(optional)</span></div>
    <div class="card-sub" style="margin-bottom:16px">Lade Dateien hoch — Bücher, Skripte, Podcasts, FAQs. Je mehr dein AI weiß, desto besser berät er.</div>

    <div id="knowledge-files-list"></div>

    <div class="file-drop-row">
      <div class="file-drop-zone" id="fdz-book" onclick="triggerFileInput('file-book')">
        <input type="file" id="file-book" accept=".pdf,.txt,.md,.docx,.html,.csv,.rtf,.json" style="display:none" onchange="handleKnowledgeFile(event,'book','📖 Buch / Skript')">
        <div class="fdz-icon">📖</div>
        <div class="fdz-label">Buch / Skript</div>
        <div class="fdz-sub">PDF · TXT · DOCX · HTML</div>
      </div>
      <div class="file-drop-zone" id="fdz-podcast" onclick="triggerFileInput('file-podcast')">
        <input type="file" id="file-podcast" accept=".pdf,.txt,.md,.docx,.html,.csv,.rtf,.json" style="display:none" onchange="handleKnowledgeFile(event,'podcast','🎙️ Podcast')">
        <div class="fdz-icon">🎙️</div>
        <div class="fdz-label">Podcast / Interview</div>
        <div class="fdz-sub">TXT · DOCX · PDF</div>
      </div>
      <div class="file-drop-zone" id="fdz-faq" onclick="triggerFileInput('file-faq')">
        <input type="file" id="file-faq" accept=".pdf,.txt,.md,.docx,.html,.csv,.rtf,.json" style="display:none" onchange="handleKnowledgeFile(event,'faq','💬 FAQ')">
        <div class="fdz-icon">💬</div>
        <div class="fdz-label">FAQ / Sales-Calls</div>
        <div class="fdz-sub">PDF · TXT · DOCX</div>
      </div>
      <div class="file-drop-zone" id="fdz-other" onclick="triggerFileInput('file-other')">
        <input type="file" id="file-other" accept=".pdf,.txt,.md,.docx,.html,.csv,.rtf,.json" style="display:none" onchange="handleKnowledgeFile(event,'other','📄 Sonstiges')">
        <div class="fdz-icon">📄</div>
        <div class="fdz-label">Sonstiges</div>
        <div class="fdz-sub">HTML · TXT · CSV · JSON</div>
      </div>
    </div>

    <div class="btn-row">
      <button class="btn btn-primary" onclick="startTranscription()">Channel transkribieren →</button>
    </div>

    <div class="pbar" id="pbar"><div class="pfill" id="pfill"></div></div>
    <div class="log-card" id="log-card">
      <div class="log-header"><div class="log-dot" id="log-dot"></div><span id="log-status">TRANSKRIBIERE...</span></div>
      <div id="log-output"></div>
    </div>

    <div class="success-box" id="success-1">
      <h3>✅ Channel erfolgreich transkribiert!</h3>
      <p>Deine Wissensbasis ist bereit. Weiter zum Profil-Setup.</p>
      <div class="btn-row" style="margin-top:16px">
        <button class="btn btn-primary" onclick="showStep(2)">Profil einrichten →</button>
      </div>
    </div>
  </div>

  <!-- STEP 2: PROFILE -->
    <!-- STEP 3: PROFILE -->
  <div class="card" id="card-2">
    <div class="card-title">👤 Creator-Profil</div>
    <div class="card-sub">Beschreibe deiner AI wer du bist, wem du hilfst und wie — das formt den Ton und die Strategie aller Gespräche.</div>

    <div class="field">
      <label>Dein Name / Brand</label>
      <input type="text" id="creator-name" placeholder="Sven Gold">
    </div>
    <div class="field">
      <label>Bio & Nische</label>
      <textarea id="creator-bio" rows="3" placeholder="Ich helfe Unternehmern ihre Sichtbarkeit durch Stimme, Kamerapräsenz und Sales-Techniken aufzubauen..."></textarea>
    </div>
    <div class="field">
      <label>Idealer Kunde (Zielgruppe)</label>
      <textarea id="creator-avatar" rows="2" placeholder="Unternehmer und Coaches die ihren Expertenstatus aufbauen wollen und high-ticket Produkte verkaufen..."></textarea>
    </div>
    <div class="field">
      <label>Booking / Call Link</label>
      <input type="text" id="booking-link" placeholder="https://calendly.com/... oder https://cal.com/...">
    </div>

    <div class="btn-row">
      <button class="btn btn-ghost" onclick="showStep(1)">← Zurück</button>
      <button class="btn btn-primary" onclick="showStep(3)">Produkte konfigurieren →</button>
    </div>
  </div>

  <!-- STEP 4: PRODUCTS -->
  <div class="card" id="card-3">
    <div class="card-title">💰 Produkte & Angebote</div>
    <div class="card-sub">Füge deine Produkte, Kurse oder Dienstleistungen hinzu. Die AI empfiehlt sie intelligent basierend auf der Situation des Besuchers.</div>

    <div class="product-list" id="product-list">
      <div class="product-item">
        <div><label>Produktname</label><input type="text" placeholder="z.B. Kommunikationstraining" class="p-name"></div>
        <div><label>URL</label><input type="text" placeholder="https://..." class="p-url"></div>
        <button class="btn-remove" onclick="removeProduct(this)">×</button>
      </div>
    </div>
    <button class="btn-add" onclick="addProduct()">+ Produkt hinzufügen</button>

    <div class="btn-row">
      <button class="btn btn-ghost" onclick="showStep(2)">← Zurück</button>
      <button class="btn btn-primary" onclick="saveAndLaunch()">🚀 Seite live schalten →</button>
    </div>
  </div>

  <!-- STEP 5: LIVE + DASHBOARD -->
  <div class="card" id="card-4">
    <div class="card-title">🎉 Deine KnowFlow Seite ist live!</div>
    <div class="card-sub">Teile diesen Link überall — dein AI-Assistent qualifiziert bereits rund um die Uhr deine Besucher.</div>

    <div class="success-box" style="display:block">
      <h3>✅ Alles konfiguriert!</h3>
      <p style="color:var(--t2);font-size:13px;margin-top:4px">Dein AI ist trainiert und bereit zu lernen, zu qualifizieren und zu verkaufen.</p>
      <div class="success-url" id="page-url" onclick="copyUrl()">Lädt...</div>
      <p style="font-size:11px;color:var(--t3);margin-top:6px;font-family:'DM Mono',monospace">Klicken zum Kopieren</p>
    </div>

    <div class="sep"></div>
    <div style="font-size:13px;font-weight:700;margin-bottom:14px;color:var(--t2)">CREATOR DASHBOARD</div>

    <div class="dash-grid">
      <div class="dash-card" onclick="openUpdateKnowledge()">
        <div class="dash-card-icon">🧠</div>
        <div class="dash-card-title">Wissen aktualisieren</div>
        <div class="dash-card-sub">Neue Quellen hinzufügen, Inhalte bearbeiten, YouTube neu scannen</div>
      </div>
      <div class="dash-card" id="view-page-card">
        <div class="dash-card-icon">👁️</div>
        <div class="dash-card-title">Seite ansehen</div>
        <div class="dash-card-sub">Deine KnowFlow Seite so sehen wie deine Besucher sie sehen</div>
      </div>
      <div class="dash-card" onclick="openEditProfile()">
        <div class="dash-card-icon">✏️</div>
        <div class="dash-card-title">Profil bearbeiten</div>
        <div class="dash-card-sub">Bio, Produkte, Booking-Link aktualisieren</div>
      </div>
      <div class="dash-card" style="opacity:.5;cursor:not-allowed">
        <div class="dash-card-icon">📊</div>
        <div class="dash-card-title">Analytics</div>
        <div class="dash-card-sub">Besucher, Leads, Conversion — coming soon</div>
      </div>
    </div>

    <!-- UPDATE KNOWLEDGE PANEL (hidden by default) -->
    <div id="update-knowledge-panel" style="display:none">
      <div class="sep"></div>
      <div class="card-title" style="font-size:16px;margin-bottom:6px">🔄 Wissen aktualisieren</div>
      <div class="card-sub" style="margin-bottom:20px">Füge neue Wissensquellen hinzu oder überschreibe bestehende. YouTube-Transkripte bleiben erhalten und werden ergänzt.</div>

      <div class="sources-list">
        <div class="source-item" id="upd-manual">
          <div class="source-header"><div class="source-type">✍️ Eigenes Wissen</div><div class="source-badge">✓ Gespeichert</div></div>
          <textarea id="upd-manual-text" class="source-textarea" rows="4"
            placeholder="Neue Erkenntnisse, aktualisierte Methoden, neue Fallstudien..."
            oninput="markSourceAdded('upd-manual', this.value)"></textarea>
        </div>
        <div class="source-item" id="upd-book">
          <div class="source-header"><div class="source-type">📖 Neuer Kursinhalt / Skript</div><div class="source-badge">✓ Gespeichert</div></div>
          <textarea id="upd-book-text" class="source-textarea" rows="4"
            placeholder="Neue Kapitel, aktualisierte Kurs-Materialien..."
            oninput="markSourceAdded('upd-book', this.value)"></textarea>
        </div>
        <div class="source-item" id="upd-faq">
          <div class="source-header"><div class="source-type">💬 Neue FAQs & Einwände</div><div class="source-badge">✓ Gespeichert</div></div>
          <textarea id="upd-faq-text" class="source-textarea" rows="4"
            placeholder="Neue häufige Fragen, aktuelle Testimonials, neue Einwand-Behandlungen..."
            oninput="markSourceAdded('upd-faq', this.value)"></textarea>
        </div>
      </div>

      <div style="margin-top:12px">
        <label style="margin-bottom:6px;display:block">YouTube-Channel neu scannen</label>
        <div style="display:flex;gap:10px;align-items:center">
          <input type="text" id="upd-channel-url" placeholder="Channel URL (leer lassen = aktuellen Channel benutzen)" style="flex:1">
          <button class="btn btn-ghost" onclick="startRescan()" style="white-space:nowrap;flex-shrink:0">🔄 Neu scannen</button>
        </div>
      </div>

      <div class="btn-row">
        <button class="btn btn-primary" onclick="saveUpdateKnowledge()">💾 Wissen speichern</button>
        <button class="btn btn-ghost" onclick="closeUpdateKnowledge()">Abbrechen</button>
      </div>
    </div>

    <div class="btn-row" style="margin-top:8px">
      <button class="btn btn-primary" id="open-page-btn">Seite öffnen →</button>
      <button class="btn btn-ghost" onclick="showStep(1)">Neuen Channel einrichten</button>
    </div>
  </div>

</div>

<script>
let currentSlug = '';
let pollInterval = null;
let logFrom = 0;
let greetingVideoFile = null; // kept for backward compat

function updateSlugPreview() {
  const val = document.getElementById('creator-slug').value.toLowerCase().replace(/[^a-z0-9-]/g, '-');
  document.getElementById('slug-display').textContent = val || 'deinname';
}

function showStep(n) {
  for (let i = 1; i <= 4; i++) {
    document.getElementById('card-'+i).classList.toggle('active', i === n);
    const step = document.getElementById('step-'+i);
    step.className = 'step' + (i === n ? ' active' : i < n ? ' done' : '');
    if (i < n) step.querySelector('.step-dot').textContent = '✓';
    else step.querySelector('.step-dot').textContent = i;
  }
  window.scrollTo({top: 0, behavior: 'smooth'});
}

// ── VIDEO UPLOAD ──
function getVideoEmbedUrl(url) {
  if (!url) return null;
  url = url.trim();
  // YouTube formats
  let yt = url.match(/(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([\w-]{11})/);
  if (yt) return 'https://www.youtube.com/embed/' + yt[1] + '?rel=0&modestbranding=1';
  // YouTube shorts
  let yts = url.match(/youtube\.com\/shorts\/([\w-]{11})/);
  if (yts) return 'https://www.youtube.com/embed/' + yts[1];
  // Vimeo
  let vm = url.match(/vimeo\.com\/(?:video\/)?([0-9]+)/);
  if (vm) return 'https://player.vimeo.com/video/' + vm[1] + '?color=ff6830&title=0&byline=0';
  return null;
}

function previewVideoUrl(url) {
  const embedUrl = getVideoEmbedUrl(url);
  const preview = document.getElementById('video-preview');
  const iframe = document.getElementById('video-iframe');
  const name = document.getElementById('video-name');
  if (embedUrl) {
    iframe.src = embedUrl;
    preview.style.display = 'block';
    name.textContent = '✅ Video erkannt — wird als Begrüßungsvideo angezeigt';
  } else if (url.length > 5) {
    name.textContent = '⚠️ Kein gültiger YouTube oder Vimeo Link';
    preview.style.display = 'none';
  } else {
    preview.style.display = 'none';
  }
}

// ── KNOWLEDGE FILE UPLOAD ──
function triggerFileInput(id) {
  document.getElementById(id).click();
}

async function handleKnowledgeFile(e, source, label) {
  const file = e.target.files[0];
  if (!file) return;
  // Resolve slug: from currentSlug, URL param, or input field
  const urlSlug = new URLSearchParams(window.location.search).get('slug') || '';
  const inputSlug = document.getElementById('creator-slug') ? document.getElementById('creator-slug').value.trim().toLowerCase().replace(/[^a-z0-9-]/g,'') : '';
  const slug = currentSlug || urlSlug || inputSlug;
  if (!slug) { alert('Bitte erst den Slug eingeben und den Channel transkribieren.'); return; }
  const zoneId = 'fdz-' + source;
  const zone = document.getElementById(zoneId);
  if (zone) { zone.style.opacity = '0.6'; zone.querySelector('.fdz-sub').textContent = 'Wird hochgeladen...'; }
  const fd = new FormData();
  fd.append('slug', slug);
  fd.append('source', source);
  fd.append('file', file);
  try {
    const r = await fetch('/api/knowledge/upload', {method:'POST', body: fd});
    const d = await r.json();
    if (d.ok) {
      if (zone) {
        zone.classList.add('done');
        zone.querySelector('.fdz-sub').textContent = '✓ ' + file.name;
        zone.style.opacity = '1';
      }
      // Add to uploaded files list
      const list = document.getElementById('knowledge-files-list');
      const item = document.createElement('div');
      item.className = 'kf-item';
      item.innerHTML = '<div class="kf-icon">' + label.split(' ')[0] + '</div>'
        + '<div class="kf-name">' + file.name + '</div>'
        + '<div class="kf-words">+' + d.words.toLocaleString('de') + ' Wörter</div>';
      list.appendChild(item);
    } else {
      if (zone) { zone.querySelector('.fdz-sub').textContent = 'Fehler: ' + (d.error||'?'); zone.style.opacity='1'; }
    }
  } catch(err) {
    if (zone) { zone.querySelector('.fdz-sub').textContent = 'Upload fehlgeschlagen'; zone.style.opacity='1'; }
  }
}

// ── CHAR COUNT ──
function updateCharCount(inputId, countId) {
  const len = document.getElementById(inputId).value.length;
  document.getElementById(countId).textContent = len.toLocaleString('de') + ' Zeichen';
}

function markSourceAdded(itemId, value) {
  document.getElementById(itemId).classList.toggle('added', value.trim().length > 10);
  updateKBPreview();
}

function updateKBPreview() {
  // file-based knowledge — count handled by kf-item list
  const filled = document.querySelectorAll('.kf-item').length;
  const kbSub = document.getElementById('kb-sub'); if(kbSub) kbSub.textContent = 'YouTube-Transkripte + ' + filled + ' Datei' + (filled !== 1 ? 'en' : '');
}

// ── TRANSCRIPTION ──
async function startTranscription() {
  const url = document.getElementById('channel-url').value.trim();
  const slug = document.getElementById('creator-slug').value.trim().toLowerCase().replace(/[^a-z0-9-]/g, '-');
  const maxV = document.getElementById('max-videos').value;
  if (!url || !slug) { alert('Bitte Channel-URL und Seiten-Namen eingeben.'); return; }
  currentSlug = slug;

  document.getElementById('log-card').style.display = 'block';
  document.getElementById('pbar').style.display = 'block';
  document.getElementById('success-1').style.display = 'none';
  document.getElementById('log-output').innerHTML = '';
  logFrom = 0;
  await fetch('/api/transcribe/start?slug='+slug+'&url='+encodeURIComponent(url)+'&max='+maxV);
  pollInterval = setInterval(pollProgress, 600);
}

async function startRescan() {
  const urlInput = document.getElementById('upd-channel-url') ? document.getElementById('upd-channel-url').value.trim() : '';
  const slug = currentSlug;
  if (!slug) { alert('Kein aktiver Creator. Bitte zuerst Setup durchführen.'); return; }
  const cfg = await fetch('/api/creator/config?slug='+slug).then(r=>r.json()).catch(()=>({}));
  const url = urlInput || cfg.channel_url || '';
  if (!url) { alert('Bitte Channel-URL eingeben.'); return; }
  // Show log area
  const logArea = document.getElementById('log-output');
  if (logArea) { logArea.innerHTML = ''; logFrom = 0; }
  const logSection = document.getElementById('log-section');
  if (logSection) logSection.style.display = 'block';
  const dot = document.getElementById('log-dot');
  if (dot) dot.className = 'log-dot running';
  const status = document.getElementById('log-status');
  if (status) status.textContent = 'LÄUFT...';
  await fetch('/api/transcribe/start?slug='+slug+'&url='+encodeURIComponent(url)+'&max=30');
  if (pollInterval) clearInterval(pollInterval);
  pollInterval = setInterval(pollProgress, 1500);
}

async function pollProgress() {
  const res = await fetch('/api/transcribe/progress?from='+logFrom);
  const data = await res.json();
  const out = document.getElementById('log-output');
  data.entries.forEach(e => {
    const div = document.createElement('div');
    div.className = 'l-'+e.type;
    div.innerHTML = e.msg;
    out.appendChild(div);
    logFrom++;
  });
  out.scrollTop = out.scrollHeight;
  const done = data.entries.filter(e => e.type === 'success').length;
  document.getElementById('pfill').style.width = Math.min(95, done * 3) + '%';
  if (data.done) {
    clearInterval(pollInterval);
    document.getElementById('log-dot').className = 'log-dot done';
    document.getElementById('log-status').textContent = 'FERTIG';
    document.getElementById('pfill').style.width = '100%';
    document.getElementById('success-1').style.display = 'block';
  }
}

// ── SAVE KNOWLEDGE ──
async function saveKnowledge() {
  // File uploads handled inline — just proceed to profile
  showStep(2);
}

async function saveUpdateKnowledge() {
  const slug = currentSlug;
  if (!slug) return;
  const sources = {
    'AKTUALISIERUNG - EIGENES WISSEN': document.getElementById('upd-manual-text').value,
    'AKTUALISIERUNG - KURSINHALT': document.getElementById('upd-book-text').value,
    'AKTUALISIERUNG - FAQ': document.getElementById('upd-faq-text').value,
  };
  const combined = Object.entries(sources)
    .filter(([k,v]) => v.trim().length > 0)
    .map(([k,v]) => '=== '+k+' ===\n'+v.trim()).join('\n\n');
  if (!combined) { alert('Bitte mindestens ein Feld ausfüllen.'); return; }
  const res = await fetch('/api/knowledge/append', {
    method: 'POST',
    headers: {'Content-Type':'application/json'},
    body: JSON.stringify({slug, text: combined, source: 'update'})
  });
  const d = await res.json();
  if (d.ok) {
    alert('✅ Wissen erfolgreich hinzugefügt!');
    closeUpdateKnowledge();
    // clear fields
    ['upd-manual-text','upd-book-text','upd-faq-text'].forEach(id => {
      document.getElementById(id).value = '';
      const item = document.getElementById(id.replace('-text',''));
      if(item) item.classList.remove('added');
    });
  }
}

// ── PRODUCTS ──
function addProduct() {
  const list = document.getElementById('product-list');
  const div = document.createElement('div');
  div.className = 'product-item';
  div.innerHTML = '<div><label>Produktname</label><input type="text" placeholder="Produktname..." class="p-name"></div>'
    +'<div><label>URL</label><input type="text" placeholder="https://..." class="p-url"></div>'
    +'<button class="btn-remove" onclick="removeProduct(this)">×</button>';
  list.appendChild(div);
}
function removeProduct(btn) {
  if (document.querySelectorAll('.product-item').length > 1) btn.parentElement.remove();
}

// ── SAVE & LAUNCH ──
async function saveAndLaunch() {
  const slug = currentSlug || document.getElementById('creator-slug').value.trim().toLowerCase().replace(/[^a-z0-9-]/g, '-');
  if (!slug) { alert('Bitte erst Channel transkribieren (Schritt 1).'); return; }
  const products = [];
  document.querySelectorAll('.product-item').forEach(item => {
    const name = item.querySelector('.p-name').value.trim();
    const url = item.querySelector('.p-url').value.trim();
    if (name) products.push({name, url});
  });
  const videoUrl = document.getElementById('greeting-video-url') ? document.getElementById('greeting-video-url').value.trim() : '';
  const body = {
    slug,
    bio: document.getElementById('creator-bio').value,
    avatar: document.getElementById('creator-avatar').value,
    booking_link: document.getElementById('booking-link').value,
    products,
    greeting_video_url: videoUrl || ''
  };
  await fetch('/api/creator/save', {method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)});
  const pageUrl = 'http://localhost:7891/c/'+slug;
  document.getElementById('page-url').textContent = pageUrl;
  document.getElementById('open-page-btn').onclick = () => window.open(pageUrl, '_blank');
  document.getElementById('view-page-card').onclick = () => window.open(pageUrl, '_blank');
  document.getElementById('dashboard-link').style.display = 'inline-flex';
  showStep(4);
}

// ── DASHBOARD ACTIONS ──
function openUpdateKnowledge() {
  document.getElementById('update-knowledge-panel').style.display = 'block';
  document.getElementById('update-knowledge-panel').scrollIntoView({behavior:'smooth'});
}
function closeUpdateKnowledge() {
  document.getElementById('update-knowledge-panel').style.display = 'none';
}
function openEditProfile() { showStep(2); }
function showDashboard() { showStep(4); }

function copyUrl() {
  const url = document.getElementById('page-url').textContent;
  navigator.clipboard.writeText(url);
  document.getElementById('page-url').textContent = '✅ Kopiert!';
  setTimeout(() => document.getElementById('page-url').textContent = url, 2000);
}

// Load existing creator if slug in URL
const urlParams = new URLSearchParams(window.location.search);
const existingSlug = urlParams.get('slug');
if (existingSlug) {
  currentSlug = existingSlug;
  document.getElementById('creator-slug').value = existingSlug;
  document.getElementById('dashboard-link').style.display = 'inline-flex';
  fetch('/api/creator/config?slug='+existingSlug).then(r=>r.json()).then(cfg => {
    if (cfg.channel_url) document.getElementById('channel-url').value = cfg.channel_url || '';
    if (cfg.bio) document.getElementById('creator-bio').value = cfg.bio;
    if (cfg.booking_link) document.getElementById('booking-link').value = cfg.booking_link;
  });
}
</script>
</body>
</html>"""




# -- VIEWER PAGE HTML --

ADMIN_HTML = r"""<!DOCTYPE html>
<html lang="de">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>KnowFlow Admin — {{SLUG}}</title>
<link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&family=DM+Mono:wght@400;500&display=swap" rel="stylesheet">
<script>window.__clerk_pk="{{CLERK_PK}}";</script>
<script async crossorigin="anonymous"
  data-clerk-publishable-key="{{CLERK_PK}}"
  src="https://cdn.jsdelivr.net/npm/@clerk/clerk-js@5/dist/clerk.browser.js">
</script>
<style>
*,*::before,*::after{box-sizing:border-box;margin:0;padding:0}
:root{
  --bg:#000;--s:#111;--s2:#1a1a1a;--b:rgba(255,255,255,.08);--b2:rgba(255,255,255,.13);
  --acc:#ff6830;--acc2:#ff9f0a;--acc-bg:rgba(255,104,48,.09);--acc-b:rgba(255,104,48,.3);
  --grad:linear-gradient(135deg,#ff6830,#ff9f0a);
  --t:#f0f0f0;--t2:#999;--t3:#555;--ok:#30d158;--err:#ff453a;--warn:#ff9f0a;
}
html,body{background:var(--bg);color:var(--t);font-family:'DM Sans',sans-serif;min-height:100vh;-webkit-font-smoothing:antialiased}
body{background-image:radial-gradient(ellipse 70% 40% at 50% 0%,rgba(255,104,48,.07) 0%,transparent 60%)}
nav{position:sticky;top:0;z-index:50;background:rgba(0,0,0,.88);backdrop-filter:blur(20px);
  border-bottom:1px solid var(--b);padding:14px 28px;display:flex;align-items:center;justify-content:space-between}
.logo{font-family:'DM Mono',monospace;font-size:14px;font-weight:500}
.logo em{background:var(--grad);-webkit-background-clip:text;-webkit-text-fill-color:transparent;font-style:normal}
.nav-r{display:flex;align-items:center;gap:8px}
.nav-slug{font-family:'DM Mono',monospace;font-size:10px;color:var(--t3);background:var(--s2);border:1px solid var(--b);border-radius:5px;padding:3px 8px}
.btn-view{background:var(--grad);color:#fff;font-size:12px;font-weight:600;padding:7px 14px;border-radius:8px;border:none;cursor:pointer;text-decoration:none;display:inline-flex;align-items:center;gap:5px}
.nav-back{color:var(--t3);font-size:12px;text-decoration:none;transition:color .2s}.nav-back:hover{color:var(--t)}
.wrap{max-width:700px;margin:0 auto;padding:28px 20px 80px}
.page-title{font-size:22px;font-weight:700;margin-bottom:6px}
.page-sub{font-size:12px;color:var(--t3);margin-bottom:20px}
.stats{display:flex;gap:12px;flex-wrap:wrap;background:var(--s2);border:1px solid var(--b);border-radius:10px;padding:14px 16px;margin-bottom:14px}
.stat-val{font-size:18px;font-weight:700;font-family:'DM Mono',monospace;background:var(--grad);-webkit-background-clip:text;-webkit-text-fill-color:transparent}
.stat-lbl{font-size:10px;color:var(--t3);text-transform:uppercase;letter-spacing:.06em;margin-top:2px}
.card{background:var(--s);border:1px solid var(--b);border-radius:14px;padding:22px;margin-bottom:12px}
.card-title{font-size:14px;font-weight:700;margin-bottom:16px;display:flex;align-items:center;gap:6px}
.badge{font-size:10px;font-family:'DM Mono',monospace;color:var(--t3);background:var(--s2);border:1px solid var(--b);border-radius:4px;padding:2px 6px;margin-left:auto}
.field{margin-bottom:14px}
label{display:block;font-size:10px;font-weight:600;color:var(--t3);letter-spacing:.06em;text-transform:uppercase;margin-bottom:5px}
input[type=text],textarea,select{width:100%;background:var(--s2);border:1px solid var(--b2);border-radius:8px;padding:9px 12px;color:var(--t);font-family:'DM Sans',sans-serif;font-size:13px;outline:none;transition:border-color .15s}
input:focus,textarea:focus{border-color:var(--acc-b)}
textarea{resize:vertical;min-height:72px;line-height:1.6}
.grid2{display:grid;grid-template-columns:1fr 1fr;gap:10px}
@media(max-width:480px){.grid2{grid-template-columns:1fr}}
.video-wrap{margin-top:10px;display:none}
.video-wrap iframe{width:100%;aspect-ratio:16/9;border-radius:8px;border:none;background:#000}
.vstatus{font-size:11px;color:var(--ok);margin-top:5px;font-family:'DM Mono',monospace}
.prod-row{display:grid;grid-template-columns:1fr 1fr 28px;gap:8px;align-items:center;margin-bottom:8px}
.btn-rm{width:26px;height:26px;border-radius:6px;border:1px solid var(--b2);background:var(--s2);color:var(--err);font-size:15px;cursor:pointer;display:flex;align-items:center;justify-content:center;flex-shrink:0}
.btn-rm:hover{background:rgba(255,69,58,.1)}
.btn-add{font-size:12px;color:var(--acc);background:none;border:1px dashed var(--acc-b);border-radius:8px;padding:7px 14px;cursor:pointer;width:100%;font-family:'DM Sans',sans-serif;font-weight:600;transition:all .15s}
.btn-add:hover{background:var(--acc-bg)}
.drop-grid{display:grid;grid-template-columns:1fr 1fr;gap:8px;margin-bottom:10px}
.dz{border:1px dashed var(--b2);border-radius:10px;padding:14px 8px;text-align:center;cursor:pointer;transition:all .2s;background:var(--s2)}
.dz:hover{border-color:var(--acc-b);background:var(--acc-bg)}
.dz.done{border-color:var(--ok);border-style:solid;background:rgba(48,209,88,.04)}
.dz.err{border-color:var(--err);border-style:solid}
.dz-ico{font-size:18px;margin-bottom:4px}
.dz-lbl{font-size:11px;font-weight:600}
.dz-sub{font-size:10px;color:var(--t3);font-family:'DM Mono',monospace;margin-top:2px}
.uploaded{display:flex;flex-direction:column;gap:5px;margin-bottom:8px}
.ui{display:flex;align-items:center;gap:8px;background:rgba(48,209,88,.05);border:1px solid rgba(48,209,88,.15);border-radius:7px;padding:6px 10px;font-size:12px}
.ui-name{flex:1}.ui-w{font-size:10px;color:var(--ok);font-family:'DM Mono',monospace}
.btn-row{display:flex;gap:8px;margin-top:16px;flex-wrap:wrap}
.btn-p{background:var(--grad);color:#fff;border:none;border-radius:9px;padding:10px 20px;font-size:13px;font-weight:700;cursor:pointer;font-family:'DM Sans',sans-serif;transition:opacity .15s}
.btn-p:hover{opacity:.82}
.btn-g{background:var(--s2);color:var(--t2);border:1px solid var(--b2);border-radius:9px;padding:10px 16px;font-size:13px;font-weight:600;cursor:pointer;font-family:'DM Sans',sans-serif;transition:all .15s}
.btn-g:hover{color:var(--t)}
.pbar{height:3px;background:var(--b);border-radius:2px;margin-top:10px;display:none}
.pfill{height:100%;background:var(--grad);width:0;transition:width .4s;border-radius:2px}
.logbox{font-family:'DM Mono',monospace;font-size:11px;color:var(--t3);background:var(--bg);border:1px solid var(--b);border-radius:8px;padding:10px;margin-top:8px;max-height:120px;overflow-y:auto;display:none;line-height:1.7}
.l-ok{color:var(--ok)}.l-err{color:var(--err)}.l-w{color:var(--warn)}
#toast{position:fixed;bottom:24px;left:50%;transform:translateX(-50%) translateY(60px);background:var(--s);border:1px solid var(--b2);border-radius:10px;padding:10px 18px;font-size:13px;font-weight:600;box-shadow:0 8px 32px rgba(0,0,0,.5);transition:transform .3s cubic-bezier(.34,1.56,.64,1);z-index:999;white-space:nowrap}
#toast.show{transform:translateX(-50%) translateY(0)}
#toast.ok{color:var(--ok);border-color:rgba(48,209,88,.25)}
#toast.err{color:var(--err);border-color:rgba(255,69,58,.25)}
</style>
</head>
<body>
<nav>
  <span class="logo">Know<em>Flow</em> <span style="color:var(--t3);font-size:11px">Admin</span></span>
  <div class="nav-r">
    <span class="nav-slug">/c/{{SLUG}}</span>
    <a href="/c/{{SLUG}}" target="_blank" class="btn-view">👁 Seite</a>
    <button id="btn-logout" onclick="doLogout()" style="display:none;background:var(--s2);color:var(--t3);border:1px solid var(--b2);border-radius:8px;padding:6px 12px;font-size:11px;font-family:'DM Mono',monospace;cursor:pointer">⬡ Logout</button>
    <a href="/" class="nav-back">← Setup</a>
  </div>
</nav>
<div class="wrap">
  <div class="page-title" id="page-title">{{SLUG}}</div>
  <div class="page-sub">Alles bearbeiten ohne neu zu transkribieren</div>

  <div class="stats" id="stats">
    <div><div class="stat-val" id="sw">—</div><div class="stat-lbl">Wörter im AI</div></div>
    <div><div class="stat-val" id="sv">—</div><div class="stat-lbl">Videos</div></div>
    <div><div class="stat-val" id="sp">—</div><div class="stat-lbl">Produkte</div></div>
    <div><div class="stat-val" id="su">—</div><div class="stat-lbl">Letztes Update</div></div>
  </div>

  <!-- PROFIL -->
  <div class="card">
    <div class="card-title">👤 Profil & Seite</div>
    <div class="grid2">
      <div class="field"><label>Name / Brand</label><input type="text" id="f-brand" placeholder="Dein Name"></div>
      <div class="field"><label>Booking Link</label><input type="text" id="f-booking" placeholder="https://calendly.com/..."></div>
    </div>
    <div class="field"><label>Bio</label><textarea id="f-bio" rows="3" placeholder="Ich helfe..."></textarea></div>
    <div class="field"><label>Zielgruppe / Idealer Kunde</label><textarea id="f-avatar" rows="2" placeholder="Unternehmer die..."></textarea></div>
    <div class="btn-row"><button class="btn-p" onclick="saveProfile()">💾 Profil speichern</button></div>
  </div>

  <!-- VIDEO -->
  <div class="card">
    <div class="card-title">🎥 Begrüßungsvideo <span class="badge">optional</span></div>
    <div class="field" style="margin-bottom:8px">
      <label>YouTube oder Vimeo URL</label>
      <input type="text" id="f-video" placeholder="https://youtube.com/watch?v=..." oninput="previewVid(this.value)">
    </div>
    <div class="video-wrap" id="vwrap">
      <iframe id="viframe" src="" allowfullscreen allow="autoplay; encrypted-media"></iframe>
      <div class="vstatus">✅ Video erkannt — wird auf deiner Seite angezeigt</div>
    </div>
    <div class="btn-row">
      <button class="btn-p" onclick="saveVideo()">💾 Video speichern</button>
      <button class="btn-g" onclick="removeVideo()">× Entfernen</button>
    </div>
  </div>

  <!-- PRODUKTE -->
  <div class="card">
    <div class="card-title">💰 Produkte & Angebote</div>
    <div id="prod-list"></div>
    <button class="btn-add" onclick="addProd()">+ Produkt hinzufügen</button>
    <div class="btn-row"><button class="btn-p" onclick="saveProds()">💾 Produkte speichern</button></div>
  </div>

  <!-- DATEIEN -->
  <div class="card">
    <div class="card-title">🧠 Wissensquellen — Dateien <span class="badge">hinzufügen</span></div>
    <p style="font-size:12px;color:var(--t3);margin-bottom:12px">Wird zur bestehenden Wissensbasis angehängt. Nichts wird überschrieben.</p>
    <div class="uploaded" id="ulist"></div>
    <div class="drop-grid">
      <div class="dz" id="dz-book" onclick="document.getElementById('fi-book').click()">
        <input type="file" id="fi-book" style="display:none" accept=".pdf,.txt,.md,.docx,.html,.csv,.rtf" onchange="uploadFile(event,'book','📖')">
        <div class="dz-ico">📖</div><div class="dz-lbl">Buch / Skript</div><div class="dz-sub">PDF · DOCX · TXT</div>
      </div>
      <div class="dz" id="dz-podcast" onclick="document.getElementById('fi-podcast').click()">
        <input type="file" id="fi-podcast" style="display:none" accept=".pdf,.txt,.md,.docx,.html,.csv,.rtf" onchange="uploadFile(event,'podcast','🎙️')">
        <div class="dz-ico">🎙️</div><div class="dz-lbl">Podcast / Interview</div><div class="dz-sub">TXT · DOCX · PDF</div>
      </div>
      <div class="dz" id="dz-faq" onclick="document.getElementById('fi-faq').click()">
        <input type="file" id="fi-faq" style="display:none" accept=".pdf,.txt,.md,.docx,.html,.csv,.rtf" onchange="uploadFile(event,'faq','💬')">
        <div class="dz-ico">💬</div><div class="dz-lbl">FAQ / Sales-Calls</div><div class="dz-sub">PDF · TXT · DOCX</div>
      </div>
      <div class="dz" id="dz-other" onclick="document.getElementById('fi-other').click()">
        <input type="file" id="fi-other" style="display:none" accept=".pdf,.txt,.md,.docx,.html,.csv,.rtf,.json" onchange="uploadFile(event,'other','📄')">
        <div class="dz-ico">📄</div><div class="dz-lbl">Sonstiges</div><div class="dz-sub">HTML · CSV · JSON</div>
      </div>
    </div>
  </div>

  <!-- TEXT WISSEN -->
  <div class="card">
    <div class="card-title">✍️ Wissen als Text</div>
    <div class="field">
      <label>Direkt Text einfügen</label>
      <textarea id="f-text" rows="5" placeholder="Neue Methoden, FAQs, Preise, Fallstudien, Testimonials..."></textarea>
    </div>
    <div class="btn-row"><button class="btn-p" onclick="saveText()">💾 Text hinzufügen</button></div>
  </div>

  <!-- RESCAN -->
  <div class="card">
    <div class="card-title">🔄 YouTube neu scannen <span class="badge">optional</span></div>
    <p style="font-size:12px;color:var(--t3);margin-bottom:12px">Neue Videos hinzufügen ohne bestehende Transkripte zu verlieren.</p>
    <div class="grid2">
      <div class="field"><label>Channel URL</label><input type="text" id="f-channel" placeholder="https://youtube.com/@..."></div>
      <div class="field"><label>Max. Videos</label>
        <select id="f-maxv"><option value="10">10 neue</option><option value="30" selected>30 neue</option><option value="50">50 neue</option><option value="0">Alle</option></select>
      </div>
    </div>
    <div class="btn-row"><button class="btn-p" onclick="rescan()">🔄 Jetzt scannen</button></div>
    <div class="pbar" id="pbar"><div class="pfill" id="pfill"></div></div>
    <div class="logbox" id="logbox"></div>
  </div>
</div>

<div id="toast"></div>

<script>
const SLUG = '{{SLUG}}';

function toast(msg, type='ok') {
  const t = document.getElementById('toast');
  t.textContent = msg; t.className = 'show ' + type;
  setTimeout(() => t.className = '', 2800);
}

function getEmbed(url) {
  if (!url) return null; url = url.trim();
  let yt = url.match(/(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([\w-]{11})/);
  if (yt) return 'https://www.youtube.com/embed/' + yt[1] + '?rel=0&modestbranding=1';
  let yts = url.match(/youtube\.com\/shorts\/([\w-]{11})/);
  if (yts) return 'https://www.youtube.com/embed/' + yts[1];
  let vm = url.match(/vimeo\.com\/(?:video\/)?([0-9]+)/);
  if (vm) return 'https://player.vimeo.com/video/' + vm[1] + '?color=ff6830&title=0&byline=0';
  return null;
}
function previewVid(url) {
  const e = getEmbed(url);
  const w = document.getElementById('vwrap');
  if (e) { document.getElementById('viframe').src = e; w.style.display = 'block'; }
  else w.style.display = 'none';
}

async function loadConfig() {
  const cfg = await fetch('/api/creator/config?slug=' + SLUG).then(r => r.json()).catch(() => ({}));
  if (cfg.channel_name) { document.getElementById('page-title').textContent = cfg.channel_name; document.getElementById('f-brand').value = cfg.channel_name; }
  if (cfg.bio) document.getElementById('f-bio').value = cfg.bio;
  if (cfg.avatar) document.getElementById('f-avatar').value = cfg.avatar;
  if (cfg.booking_link) document.getElementById('f-booking').value = cfg.booking_link;
  if (cfg.channel_url) document.getElementById('f-channel').value = cfg.channel_url;
  if (cfg.greeting_video_url) { document.getElementById('f-video').value = cfg.greeting_video_url; previewVid(cfg.greeting_video_url); }
  document.getElementById('sw').textContent = cfg.word_count ? cfg.word_count.toLocaleString('de') : '—';
  document.getElementById('sv').textContent = cfg.video_count || '—';
  document.getElementById('sp').textContent = (cfg.products || []).filter(p => p.name).length || '—';
  document.getElementById('su').textContent = cfg.last_knowledge_update || '—';
  renderProds(cfg.products || []);
}

function renderProds(prods) {
  const el = document.getElementById('prod-list');
  el.innerHTML = '';
  (prods.length ? prods : [{}]).forEach(p => {
    const row = document.createElement('div'); row.className = 'prod-row';
    row.innerHTML = '<input type="text" class="pn" placeholder="Produktname" value="'+(p.name||'')+'">'
      + '<input type="text" class="pu" placeholder="https://..." value="'+(p.url||'')+'">'
      + '<button class="btn-rm" onclick="this.parentElement.remove()">×</button>';
    el.appendChild(row);
  });
}
function addProd() {
  const el = document.getElementById('prod-list');
  const row = document.createElement('div'); row.className = 'prod-row';
  row.innerHTML = '<input type="text" class="pn" placeholder="Produktname"><input type="text" class="pu" placeholder="https://..."><button class="btn-rm" onclick="this.parentElement.remove()">×</button>';
  el.appendChild(row); row.querySelector('.pn').focus();
}
function getProds() {
  return [...document.querySelectorAll('#prod-list .prod-row')].map(r => ({name:r.querySelector('.pn').value.trim(),url:r.querySelector('.pu').value.trim()})).filter(p=>p.name);
}

function baseBody() {
  const prods = getProds();
  const body = { slug: SLUG };
  // Only send products if there are any — prevents accidental clearing
  // If user explicitly removed all products, they need to save with the products card button
  if (prods.length > 0) body.products = prods;
  
  const fields = {
    channel_name: 'f-brand',
    bio: 'f-bio',
    avatar: 'f-avatar',
    booking_link: 'f-booking',
    greeting_video_url: 'f-video',
    channel_url: 'f-channel'
  };
  for (const [key, id] of Object.entries(fields)) {
    const el = document.getElementById(id);
    if (el && el.value.trim()) body[key] = el.value.trim();
  }
  if (!body.greeting_video_url) body.greeting_video_url = '';
  return body;
}

// Dedicated product save — explicitly clears if empty
async function saveProdsExplicit() {
  const prods = getProds();
  const body = { slug: SLUG, products: prods, clear_products: true,
    greeting_video_url: document.getElementById('f-video').value.trim() };
  ['channel_name','bio','avatar','booking_link','channel_url'].forEach(k => {
    const id = {channel_name:'f-brand',bio:'f-bio',avatar:'f-avatar',booking_link:'f-booking',channel_url:'f-channel'}[k];
    const el = document.getElementById(id);
    if (el && el.value.trim()) body[k] = el.value.trim();
  });
  const r = await fetch('/api/creator/save',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)}).then(r=>r.json()).catch(()=>null);
  document.getElementById('sp').textContent = prods.length;
  if (r && r.ok) toast('✅ Produkte gespeichert'); else toast('❌ Fehler','err');
}

async function post(body) {
  const r = await fetch('/api/creator/save',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(body)}).then(r=>r.json()).catch(()=>null);
  return r && r.ok;
}

async function saveProfile() { if(await post(baseBody())) toast('✅ Profil gespeichert'); else toast('❌ Fehler','err'); }
async function saveVideo() { if(await post(baseBody())) toast('✅ Video gespeichert'); else toast('❌ Fehler','err'); }
async function removeVideo() { document.getElementById('f-video').value=''; document.getElementById('vwrap').style.display='none'; document.getElementById('viframe').src=''; await post(baseBody()); toast('Video entfernt'); }
async function saveProds() { await saveProdsExplicit(); }

async function uploadFile(e, source, icon) {
  const file = e.target.files[0]; if (!file) return;
  const zone = document.getElementById('dz-' + source);
  zone.querySelector('.dz-sub').textContent = '⏳ Verarbeite...'; zone.style.opacity = '.6';
  const fd = new FormData(); fd.append('slug', SLUG); fd.append('source', source); fd.append('file', file);
  try {
    const d = await fetch('/api/knowledge/upload',{method:'POST',body:fd}).then(r=>r.json());
    zone.style.opacity = '1';
    if (d.ok) {
      zone.classList.add('done');
      zone.querySelector('.dz-lbl').textContent = '✅ ' + file.name.slice(0,16);
      zone.querySelector('.dz-sub').textContent = '+' + d.words.toLocaleString('de') + ' Wörter';
      const ul = document.getElementById('ulist');
      const it = document.createElement('div'); it.className = 'ui';
      it.innerHTML = '<span>'+icon+'</span><span class="ui-name">'+file.name+'</span><span class="ui-w">+'+d.words.toLocaleString('de')+' Wörter</span>';
      ul.appendChild(it);
      const sw = document.getElementById('sw'); const cur = parseInt(sw.textContent.replace(/\D/g,''))||0; sw.textContent=(cur+d.words).toLocaleString('de');
      toast('✅ '+file.name+' hinzugefügt');
    } else {
      zone.classList.add('err');
      zone.querySelector('.dz-sub').textContent = '❌ ' + (d.error||'?').slice(0,50);
      toast('❌ ' + (d.error||'Upload fehlgeschlagen'), 'err');
    }
  } catch(err) { zone.style.opacity='1'; zone.querySelector('.dz-sub').textContent='❌ Netzwerkfehler'; toast('❌ Netzwerkfehler','err'); }
  e.target.value = '';
}

async function saveText() {
  const text = document.getElementById('f-text').value.trim();
  if (!text) { toast('Bitte Text eingeben','err'); return; }
  const r = await fetch('/api/knowledge/append',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({slug:SLUG,text,source:'admin_manual'})}).then(r=>r.json()).catch(()=>null);
  if (r && r.ok) { document.getElementById('f-text').value=''; toast('✅ Text hinzugefügt'); } else toast('❌ Fehler','err');
}

let scanPoll;
async function rescan() {
  const url = document.getElementById('f-channel').value.trim();
  if (!url) { toast('Bitte Channel URL eingeben','err'); return; }
  const max = document.getElementById('f-maxv').value;
  document.getElementById('pbar').style.display='block';
  document.getElementById('logbox').style.display='block';
  document.getElementById('logbox').innerHTML='';
  document.getElementById('pfill').style.width='5%';
  await fetch('/api/transcribe/start?slug='+SLUG+'&url='+encodeURIComponent(url)+'&max='+max);
  let from=0;
  scanPoll = setInterval(async()=>{
    const lg = await fetch('/api/transcribe/progress?from='+from).then(r=>r.json()).catch(()=>({lines:[]}));
    const box = document.getElementById('logbox');
    (lg.lines||[]).forEach(l=>{ const d=document.createElement('div'); d.innerHTML=l.replace(/^✅/,'<span class=l-ok>✅</span>').replace(/^⚠️/,'<span class=l-w>⚠️</span>').replace(/^❌/,'<span class=l-err>❌</span>'); box.appendChild(d); from++; });
    box.scrollTop=box.scrollHeight;
    document.getElementById('pfill').style.width=Math.min(95,5+from*2)+'%';
    if (lg.done) { clearInterval(scanPoll); document.getElementById('pfill').style.width='100%'; toast('✅ Scan abgeschlossen'); loadConfig(); }
  },1000);
}

// ── Clerk Auth Guard ──────────────────────────────────────────────────────────
async function initAuth() {
  const pk = window.__clerk_pk;
  if (!pk || !pk.startsWith('pk_')) {
    console.warn('No Clerk key — running without auth');
    document.getElementById('btn-logout').style.display = 'none';
    loadConfig();
    return;
  }
  try {
    await window.Clerk.load();
    const user = window.Clerk.user;
    if (!user) {
      window.location.href = '/login';
      return;
    }
    // Show logout button
    const email = user.primaryEmailAddress?.emailAddress || '';
    const btn = document.getElementById('btn-logout');
    btn.style.display = 'flex';
    btn.title = email;
    const adminEmail = 'sven.gold.official@gmail.com';
    const isAdmin = email === adminEmail;

    if (isAdmin) {
      loadConfig();
      return;
    }

    // Check subscription for non-admins
    const token = await window.Clerk.session.getToken().catch(()=>'');
    const userId = window.Clerk.user?.id || '';
    const authUrl = `/api/auth/me?user_id=${encodeURIComponent(userId)}&email=${encodeURIComponent(email)}`;
    const res = await fetch(authUrl, {
      headers: token ? { 'Authorization': 'Bearer ' + token } : {}
    });
    const data = await res.json();
    const creator = data.creator || {};
    const status = creator.subscription_status || 'inactive';

    if (status !== 'active') {
      showPaywall(email, creator.slug || SLUG);
      return;
    }

    loadConfig();
  } catch(e) {
    console.warn('Auth check error:', e);
    // Don't silently allow access — show paywall on error
    const emailFallback = window.Clerk?.user?.primaryEmailAddress?.emailAddress || '';
    const adminEmail = 'sven.gold.official@gmail.com';
    if (emailFallback === adminEmail) {
      loadConfig();
    } else {
      showPaywall(emailFallback, SLUG);
    }
  }
}

async function showPaywall(email, slug) {
  document.querySelector('.wrap').innerHTML = `
    <div style="max-width:480px;margin:60px auto;text-align:center">
      <div style="font-size:40px;margin-bottom:16px">🔒</div>
      <h2 style="font-size:24px;font-weight:700;margin-bottom:8px">KnowFlow freischalten</h2>
      <p style="color:var(--t3);font-size:14px;margin-bottom:32px;line-height:1.6">
        Dein AI Sales Funnel ist einen Klick entfernt.<br>
        <strong style="color:var(--t)">29€/Monat</strong> — jederzeit kündbar.
      </p>
      <div style="background:var(--s);border:1px solid var(--b);border-radius:16px;padding:28px;margin-bottom:20px">
        <div style="display:flex;flex-direction:column;gap:10px;margin-bottom:24px;text-align:left">
          ${['✅ Unlimitierter AI-Chat für deine Viewers','✅ Automatischer Sales Funnel','✅ Lead-Tracking & Analytics','✅ Deine eigene KnowFlow-Seite','✅ Kein KnowFlow-Branding'].map(f=>`<div style="font-size:13px;color:var(--t2)">${f}</div>`).join('')}
        </div>
        <button onclick="startCheckout('${email}','${slug}')"
          style="width:100%;background:linear-gradient(135deg,#ff6830,#ff9f0a);color:#fff;border:none;border-radius:12px;padding:14px 24px;font-size:15px;font-weight:700;cursor:pointer;font-family:'DM Sans',sans-serif">
          Jetzt starten — 29€/Monat →
        </button>
      </div>
      <p style="font-size:11px;color:var(--t3)">Sicher bezahlen mit Stripe · Jederzeit kündbar</p>
    </div>`;
}

async function startCheckout(email, slug) {
  const btn = event.target;
  btn.textContent = 'Weiterleitung...';
  btn.disabled = true;
  try {
    const baseUrl = window.location.origin;
    const res = await fetch('/api/stripe/checkout', {
      method: 'POST',
      headers: {'Content-Type':'application/json'},
      body: JSON.stringify({
        slug: slug,
        plan: 'creator',
        email: email,
        clerk_user_id: window.Clerk.user?.id || '',
        base_url: baseUrl,
        success_slug: slug
      })
    });
    const d = await res.json();
    if (d.url) {
      window.location.href = d.url;
    } else {
      btn.textContent = 'Fehler: ' + (d.error || 'Unbekannt');
      btn.disabled = false;
    }
  } catch(e) {
    btn.textContent = 'Fehler — bitte neu versuchen';
    btn.disabled = false;
  }
}

async function doLogout() {
  try {
    await window.Clerk.signOut();
  } catch(e) {}
  window.location.href = '/login';
}

// Handle payment success redirect
const urlParams = new URLSearchParams(window.location.search);
if (urlParams.get('payment') === 'success') {
  // Clean up URL without reload
  const cleanUrl = window.location.pathname + '?slug=' + SLUG;
  window.history.replaceState({}, '', cleanUrl);
}

// Start auth check then load
if (window.__clerk_pk && window.__clerk_pk.startsWith('pk_')) {
  window.addEventListener('load', initAuth);
} else {
  loadConfig();
}
</script>
</body>
</html>"""


VIEWER_HTML = r"""<!DOCTYPE html>
<html lang="de">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1,maximum-scale=1,viewport-fit=cover">
<meta name="apple-mobile-web-app-capable" content="yes">
<title>{{CREATOR_NAME}}</title>
<link href="https://fonts.googleapis.com/css2?family=Instrument+Sans:ital,wght@0,400;0,500;0,600;0,700;1,400&family=Instrument+Serif:ital@0;1&display=swap" rel="stylesheet">
<script src="https://cdnjs.cloudflare.com/ajax/libs/marked/9.1.6/marked.min.js"></script>
<style>
*,::before,::after{box-sizing:border-box;margin:0;padding:0;-webkit-tap-highlight-color:transparent}
:root{
  --bg:#0d0d0f;
  --s1:rgba(255,255,255,.03);
  --s2:rgba(255,255,255,.05);
  --s3:rgba(255,255,255,.07);
  --s4:rgba(255,255,255,.10);
  --border:rgba(255,255,255,.08);
  --border2:rgba(255,255,255,.13);
  --t1:#f0f0f5;--t2:rgba(240,240,245,.55);--t3:rgba(240,240,245,.3);
  --acc:#ff6830;--acc2:#ff9f0a;
  --acc-dim:rgba(255,104,48,.10);--acc-brd:rgba(255,104,48,.30);
  --acc-txt:#ff8c55;
  --grad:linear-gradient(90deg,#ff6830,#ff9f0a);
  --ok:#30d158;--ok-dim:rgba(48,209,88,.1);--ok-brd:rgba(48,209,88,.25);
  --err:#ff453a;--err-dim:rgba(255,69,58,.09);--err-brd:rgba(255,69,58,.25);
  --glass:rgba(28,28,35,.72);
  --glass-brd:rgba(255,255,255,.1);
  --blur:blur(24px) saturate(180%);
  --safe-b:env(safe-area-inset-bottom,0px);
  --safe-t:env(safe-area-inset-top,0px);
  --r:14px;
}
html{background:var(--bg);-webkit-font-smoothing:antialiased;color-scheme:dark}
body{
  font-family:'Instrument Sans',system-ui,sans-serif;
  background:var(--bg);color:var(--t1);
  min-height:100dvh;display:flex;flex-direction:column;
  overflow-x:hidden;
  background-image:
    radial-gradient(ellipse 80% 50% at 50% -10%, rgba(120,100,200,.12) 0%, transparent 60%),
    radial-gradient(ellipse 60% 40% at 80% 80%, rgba(60,80,180,.06) 0%, transparent 50%);
}
::-webkit-scrollbar{width:2px}
::-webkit-scrollbar-thumb{background:var(--s4)}

/* ════════════ HEADER ════════════ */
#hdr{
  position:sticky;top:0;z-index:200;
  background:rgba(13,13,18,.8);
  backdrop-filter:var(--blur);-webkit-backdrop-filter:var(--blur);
  border-bottom:1px solid var(--glass-brd);
  padding:calc(10px + var(--safe-t)) 18px 10px;
  display:flex;align-items:center;gap:11px;
}
#hdr-av{
  width:36px;height:36px;border-radius:50%;flex-shrink:0;
  background:var(--acc);
  display:flex;align-items:center;justify-content:center;
  color:#000;font-weight:700;font-size:14px;letter-spacing:-.01em;
}
#hdr-name{font-size:14px;font-weight:600;letter-spacing:-.02em;flex:1;min-width:0;overflow:hidden;text-overflow:ellipsis;white-space:nowrap}
#hdr-tag{
  font-size:9px;letter-spacing:.1em;text-transform:uppercase;
  color:var(--t3);display:flex;align-items:center;gap:4px;flex-shrink:0;
}
.live-dot{
  width:5px;height:5px;border-radius:50%;background:var(--acc);
  animation:live 2s ease-in-out infinite;
}
@keyframes live{0%,100%{opacity:1}50%{opacity:.3}}

/* ════════════ PAGE SCROLL ════════════ */
#page{
  flex:1;overflow-y:auto;overflow-x:hidden;
  padding-bottom:calc(var(--safe-b) + 96px);
  scroll-behavior:smooth;
}

/* ════════════ GREETING VIDEO ════════════ */
#video-sec{display:none;padding:14px 14px 0}  /* hidden until loadConfig */
#video-sec video{
  width:100%;aspect-ratio:16/9;border-radius:var(--r);
  background:#000;display:block;object-fit:cover;
}

/* ════════════ PROFILE BLOCK ════════════ */
#profile-sec{padding:14px 18px 0}
#bio-text{
  font-size:15px;color:var(--t1);line-height:1.6;
  font-weight:600;letter-spacing:-.02em;
}
#tags-row{
  display:flex;gap:5px;margin-top:11px;
  overflow-x:auto;scrollbar-width:none;
}
#tags-row::-webkit-scrollbar{display:none}
.tag{
  flex-shrink:0;
  border:1px solid var(--border2);border-radius:20px;
  padding:3px 10px;font-size:10px;color:var(--t3);
  font-family:'Instrument Sans',sans-serif;letter-spacing:.02em;
  font-weight:500;
}
.tag.accent{border-color:var(--acc-brd);color:var(--acc-txt)}

/* ════════════ SECTION LABEL ════════════ */
.sec-label{
  padding:14px 18px 6px;
  font-size:9px;letter-spacing:.12em;text-transform:uppercase;
  color:var(--t3);font-weight:600;
}

/* ════════════ ACTION BUTTONS (Links + Products + Booking) ════════════ */
#links-sec{padding:0 14px}
#products-sec{padding:0 14px;display:none}
#booking-sec{padding:0 14px}
#booking-lnk{display:none}

/* unified action grid — all buttons same size, same style */
.act-grid{
  display:grid;gap:6px;
  grid-template-columns:repeat(3,1fr);
}
.act-grid.has-2{grid-template-columns:repeat(2,1fr)}
.act-grid.has-4{grid-template-columns:repeat(4,1fr)}
.act-grid.has-1{grid-template-columns:1fr}

.act-btn{
  display:flex;flex-direction:column;align-items:center;justify-content:center;
  gap:8px;text-align:center;
  background:var(--glass);
  backdrop-filter:var(--blur);-webkit-backdrop-filter:var(--blur);
  border:1px solid var(--glass-brd);
  border-radius:14px;padding:14px 8px;
  text-decoration:none;color:var(--t1);cursor:pointer;
  transition:all .15s cubic-bezier(.34,1.2,.64,1);
  min-height:82px;
  box-shadow:0 4px 16px rgba(0,0,0,.3), inset 0 1px 0 rgba(255,255,255,.06);
}
.act-btn:active{transform:scale(.97);background:var(--s4)}
.act-btn.booking{
  border-color:var(--acc-brd);
  background:rgba(255,104,48,.10);
  box-shadow:0 4px 16px rgba(0,0,0,.3), inset 0 1px 0 rgba(255,104,48,.12);
}
.act-btn.booking:active{transform:scale(.97)}
.act-icon{
  width:32px;height:32px;border-radius:9px;
  background:rgba(255,255,255,.08);
  border:1px solid rgba(255,255,255,.1);
  display:flex;align-items:center;justify-content:center;
  flex-shrink:0;
}
.act-btn.booking .act-icon{background:rgba(255,104,48,.12);border-color:var(--acc-brd)}
.act-icon svg{width:15px;height:15px;fill:var(--t2)}
.act-btn.booking .act-icon svg{fill:var(--acc-txt)}
.act-label{font-size:11px;font-weight:700;letter-spacing:-.01em;line-height:1.2}
.act-sub{font-size:9px;color:var(--t3);margin-top:1px}
.act-btn.booking .act-label{color:var(--acc-txt)}

/* ════════════ DIVIDER ════════════ */
.div{height:1px;background:var(--border);margin:20px 18px 0}

/* ════════════ XP STRIP ════════════ */
#xp-sec{padding:16px 18px 0;display:flex;align-items:center;gap:8px}
.xp-pill{
  border:1px solid rgba(255,255,255,.1);border-radius:20px;
  padding:4px 11px;font-size:10px;font-weight:600;color:var(--t2);
  letter-spacing:.02em;white-space:nowrap;
  background:rgba(255,255,255,.04);
  backdrop-filter:blur(8px);-webkit-backdrop-filter:blur(8px);
}
.xp-pill.streak-pill{
  border-color:rgba(255,104,48,.28);color:var(--acc-txt);
  background:var(--acc-dim);
}
.xp-pill.streak-pill.hot{
  border-color:rgba(255,159,10,.25);color:#ff9f0a;
  background:rgba(255,159,10,.07);
}
#xp-label{font-size:10px;color:var(--t3);margin-left:auto;font-weight:500}
#xp-track{height:2px;background:var(--border);border-radius:1px;margin:8px 18px 0;overflow:hidden}
#xp-fill{height:100%;background:var(--grad);border-radius:1px;width:0%;transition:width .5s cubic-bezier(.34,1.56,.64,1)}

/* ════════════ AI SECTION ════════════ */
#ai-sec{padding:20px 18px 0}
#ai-hdr{display:flex;align-items:center;gap:10px;margin-bottom:14px}
#ai-icon{
  width:32px;height:32px;border-radius:9px;flex-shrink:0;
  background:var(--acc-dim);border:1px solid var(--acc-brd);
  display:flex;align-items:center;justify-content:center;
}
#ai-icon svg{width:15px;height:15px;fill:var(--acc-txt)}
#ai-label{font-size:13px;font-weight:700;letter-spacing:-.01em}
#ai-sublabel{font-size:10px;color:var(--t3);margin-top:1px;letter-spacing:.01em}

/* Quick actions */
#qa-row{display:flex;gap:6px;flex-wrap:wrap;margin-bottom:12px}
.qa{
  display:inline-flex;align-items:center;gap:5px;
  border:1px solid rgba(255,255,255,.1);border-radius:20px;
  padding:7px 13px;font-size:11px;font-weight:600;
  color:var(--t2);cursor:pointer;
  background:rgba(255,255,255,.05);
  backdrop-filter:blur(12px);-webkit-backdrop-filter:blur(12px);
  transition:all .15s cubic-bezier(.34,1.2,.64,1);
  -webkit-user-select:none;user-select:none;
  box-shadow:0 2px 8px rgba(0,0,0,.2), inset 0 1px 0 rgba(255,255,255,.06);
}
.qa:active,.qa:hover{
  border-color:rgba(255,255,255,.2);color:var(--t1);
  background:rgba(255,255,255,.1);transform:scale(1.03);
}

/* Messages */
#msg-list{display:flex;flex-direction:column;gap:10px;margin-bottom:10px}
.msg{animation:fadeUp .18s ease both}
@keyframes fadeUp{from{opacity:0;transform:translateY(4px)}to{opacity:1;transform:none}}

.msg-u{align-self:flex-end;max-width:80%}
.msg-u .bbl{
  background:var(--grad);color:#fff;
  border-radius:16px 16px 3px 16px;
  padding:9px 13px;font-size:13px;font-weight:500;line-height:1.55;
}

.msg-a{align-self:flex-start;width:100%}
.msg-a .bbl{
  background:var(--s2);border:1px solid var(--border);
  border-radius:3px 16px 16px 16px;
  padding:11px 13px;font-size:13px;color:var(--t1);line-height:1.7;
}
.msg-a .bbl p{margin-bottom:6px}.msg-a .bbl p:last-child{margin:0}
.msg-a .bbl strong{font-weight:600;color:var(--t1)}
.msg-a .bbl ul,.msg-a .bbl ol{padding-left:16px;margin:5px 0}
.msg-a .bbl li{margin-bottom:3px}
.msg-lbl{font-size:9px;color:var(--t3);margin-top:3px;padding:0 2px;letter-spacing:.03em;font-weight:500}
.msg-u .msg-lbl{text-align:right}

/* Typing */
.typing-bbl{
  background:var(--s2);border:1px solid var(--border);
  border-radius:3px 16px 16px 16px;
  padding:12px 14px;width:54px;
  display:flex;gap:4px;align-items:center;
}
.typing-bbl span{
  width:4px;height:4px;border-radius:50%;background:var(--t3);
  animation:blink 1.3s ease-in-out infinite;
}
.typing-bbl span:nth-child(2){animation-delay:.2s}
.typing-bbl span:nth-child(3){animation-delay:.4s}
@keyframes blink{0%,80%,100%{opacity:.15;transform:scale(1)}40%{opacity:1;transform:scale(1.4)}}

/* Answer options */
.ans-opts{display:flex;flex-direction:column;gap:6px;margin-top:8px}
.ans-o{
  display:flex;align-items:center;gap:10px;
  background:var(--s2);border:1px solid var(--border);
  border-radius:10px;padding:10px 12px;
  cursor:pointer;font-size:13px;font-weight:500;color:var(--t1);
  text-align:left;-webkit-user-select:none;
  transition:border-color .12s,background .12s,color .12s;
}
.ans-o:active,.ans-o:hover{border-color:var(--acc-brd);background:var(--acc-dim);color:var(--acc-txt)}
.ans-bullet{
  width:6px;height:6px;border-radius:50%;
  border:1.5px solid var(--t3);flex-shrink:0;
  transition:all .12s;
}
.ans-o:hover .ans-bullet,.ans-o:active .ans-bullet{border-color:var(--acc);background:var(--acc)}

/* Quiz */
.qz{
  background:var(--glass);
  backdrop-filter:var(--blur);-webkit-backdrop-filter:var(--blur);
  border:1px solid var(--glass-brd);
  border-radius:18px;padding:16px;
  animation:fadeUp .25s cubic-bezier(.34,1.2,.64,1);
  box-shadow:0 8px 32px rgba(0,0,0,.4), inset 0 1px 0 rgba(255,255,255,.07);
}
.qz-tag{
  display:inline-flex;align-items:center;gap:5px;
  background:rgba(255,255,255,.07);border:1px solid rgba(255,255,255,.12);
  border-radius:20px;padding:3px 10px;
  font-size:9px;font-weight:700;color:var(--t2);
  letter-spacing:.1em;text-transform:uppercase;margin-bottom:12px;
}
.qz-diff-tag{
  margin-left:3px;font-size:8px;color:var(--t3);
  background:rgba(255,255,255,.05);
  border:1px solid rgba(255,255,255,.08);border-radius:20px;
  padding:1px 6px;letter-spacing:.04em;
}
.qz-q{
  font-size:14px;font-weight:600;line-height:1.55;
  margin-bottom:13px;letter-spacing:-.02em;color:var(--t1);
}
.qz-opts{display:flex;flex-direction:column;gap:7px}
.qz-o{
  display:flex;align-items:center;gap:11px;
  background:rgba(255,255,255,.04);
  border:1px solid rgba(255,255,255,.09);
  border-radius:12px;padding:11px 13px;
  cursor:pointer;font-size:13px;font-weight:500;color:var(--t1);
  text-align:left;-webkit-user-select:none;
  transition:all .15s cubic-bezier(.34,1.2,.64,1);
}
.qz-o:active:not(:disabled),.qz-o:hover:not(:disabled){
  background:rgba(255,255,255,.08);
  border-color:rgba(255,255,255,.18);
  transform:scale(1.01);
}
.qz-o:disabled{cursor:default;transform:none!important}
.qz-o.correct{
  border-color:var(--ok-brd)!important;
  background:var(--ok-dim)!important;color:var(--ok)!important;
  box-shadow:0 0 0 1px var(--ok-brd), 0 4px 12px rgba(48,209,88,.12)!important;
}
.qz-o.wrong{
  border-color:var(--err-brd)!important;
  background:var(--err-dim)!important;color:var(--err)!important;
  box-shadow:0 0 0 1px var(--err-brd)!important;
}
.qz-k{
  width:22px;height:22px;min-width:22px;border-radius:7px;flex-shrink:0;
  border:1px solid rgba(255,255,255,.12);
  background:rgba(255,255,255,.06);
  display:flex;align-items:center;justify-content:center;
  font-size:10px;font-weight:700;color:var(--t3);letter-spacing:.02em;
}
.qz-o.correct .qz-k{border-color:var(--ok);color:var(--ok);background:var(--ok-dim)}
.qz-o.wrong .qz-k{border-color:var(--err);color:var(--err);background:var(--err-dim)}
.qz-fb{
  margin-top:11px;padding:10px 13px;border-radius:11px;
  font-size:13px;font-weight:600;display:none;line-height:1.5;
  display:flex;align-items:center;gap:8px;
}
.qz-fb{display:none}
.qz-fb.show{display:flex}
.qz-fb-icon{font-size:16px;flex-shrink:0}
.qz-fb.ok{
  background:rgba(48,209,88,.1);color:var(--ok);
  border:1px solid var(--ok-brd);
  box-shadow:0 0 20px rgba(48,209,88,.08);
}
.qz-fb.fail{
  background:rgba(255,69,58,.09);color:var(--err);
  border:1px solid var(--err-brd);
}
.qz-exp{
  margin-top:8px;padding:11px 13px;border-radius:11px;
  font-size:12px;color:var(--t2);line-height:1.65;
  background:rgba(255,255,255,.04);
  border:1px solid rgba(255,255,255,.08);
  display:none;
}
.qz-exp.show{display:block;animation:fadeUp .2s ease}

/* Warm lead */
.warm-card{
  background:var(--glass);
  backdrop-filter:var(--blur);-webkit-backdrop-filter:var(--blur);
  border:1px solid rgba(255,104,48,.28);border-radius:18px;padding:16px;
  animation:fadeUp .25s cubic-bezier(.34,1.2,.64,1);
  box-shadow:0 8px 32px rgba(0,0,0,.4), inset 0 1px 0 rgba(255,104,48,.10);
}
.warm-card-title{font-size:12px;font-weight:700;color:var(--acc-txt);margin-bottom:5px;letter-spacing:.01em}
.warm-card-body{font-size:12px;color:var(--t2);line-height:1.6;margin-bottom:11px}
.warm-btn{
  display:inline-flex;align-items:center;gap:6px;
  background:var(--acc);color:#000;
  padding:9px 15px;border-radius:9px;
  text-decoration:none;font-size:12px;font-weight:700;
  transition:transform .12s;
}
.warm-btn:active{transform:scale(.97)}
.warm-btn svg{width:12px;height:12px;fill:#05111f}

/* ════════════ TOPIC CHIPS ════════════ */
.topic-chips{
  padding:0 16px 8px;display:flex;flex-wrap:wrap;gap:7px;
  animation:fadeUp .2s ease;
}
.topic-chip{
  display:inline-flex;align-items:center;
  background:rgba(255,104,48,.1);
  border:1px solid rgba(255,104,48,.28);
  border-radius:20px;padding:8px 15px;
  font-size:12px;font-weight:600;color:var(--acc-txt);
  cursor:pointer;-webkit-user-select:none;
  backdrop-filter:blur(12px);-webkit-backdrop-filter:blur(12px);
  transition:all .15s cubic-bezier(.34,1.2,.64,1);
  font-family:inherit;
}
.topic-chip:active,.topic-chip:hover{
  background:rgba(255,104,48,.15);
  border-color:rgba(255,104,48,.4);
  color:#e8f5ff;transform:scale(1.04);
}

/* ════════════ VOICE WAVEFORM ════════════ */
#voice-wave{
  display:flex;align-items:center;gap:10px;flex:1;
  height:42px;padding:0 4px;
}
#voice-cancel{
  width:28px;height:28px;border-radius:50%;flex-shrink:0;
  background:rgba(255,69,58,.15);border:1px solid rgba(255,69,58,.3);
  color:var(--err);cursor:pointer;
  display:flex;align-items:center;justify-content:center;
  transition:background .12s;
}
#voice-cancel:active{background:rgba(255,69,58,.25)}
.wave-bars{
  display:flex;align-items:center;gap:2px;flex:1;height:28px;
}
.wave-bar{
  flex:1;border-radius:2px;min-height:3px;max-height:24px;
  background:var(--acc);
  transition:height .08s ease;
  opacity:.8;
}
#voice-dur{
  font-size:11px;font-weight:600;color:var(--t2);
  font-variant-numeric:tabular-nums;white-space:nowrap;flex-shrink:0;
}

/* Powered by */
#pwrd{
  display:flex;justify-content:center;align-items:center;gap:5px;
  padding:24px 18px 10px;font-size:9px;color:var(--t3);
  text-decoration:none;letter-spacing:.08em;text-transform:uppercase;
  font-weight:600;
  transition:color .15s;
}
#pwrd:hover{color:var(--t2)}
#pwrd-dot{width:4px;height:4px;border-radius:50%;background:var(--acc)}

/* ════════════ FIXED INPUT ════════════ */
#inp-wrap{
  position:fixed;bottom:0;left:0;right:0;z-index:100;
  background:rgba(10,10,10,.95);
  backdrop-filter:blur(32px);-webkit-backdrop-filter:blur(32px);
  border-top:1px solid var(--border);
  padding:9px 14px calc(9px + var(--safe-b));
}
#inp-actions{
  display:flex;gap:5px;margin-bottom:8px;
  overflow-x:auto;scrollbar-width:none;
}
#inp-actions::-webkit-scrollbar{display:none}
.ia{
  display:inline-flex;align-items:center;gap:4px;flex-shrink:0;
  border:1px solid rgba(255,255,255,.1);border-radius:20px;
  padding:6px 12px;font-size:10px;font-weight:600;
  color:var(--t2);cursor:pointer;
  background:rgba(255,255,255,.05);
  backdrop-filter:blur(12px);-webkit-backdrop-filter:blur(12px);
  transition:all .15s;-webkit-user-select:none;letter-spacing:.01em;
}
.ia:active,.ia:hover{border-color:rgba(255,255,255,.2);color:var(--t1);background:rgba(255,255,255,.09)}
.ia.active-quiz{border-color:rgba(255,255,255,.18);color:var(--t1);background:rgba(255,255,255,.09)}
.ia svg{width:10px;height:10px;fill:currentColor}

#inp-row{
  display:flex;align-items:flex-end;gap:7px;
  background:var(--s2);border:1px solid var(--border);
  border-radius:20px;padding:6px 6px 6px 14px;
  transition:border-color .18s;
}
#inp-row:focus-within{border-color:var(--acc-brd)}
#chat-input{
  flex:1;background:none;border:none;outline:none;
  color:var(--t1);font-family:'Instrument Sans',sans-serif;
  font-size:14px;resize:none;max-height:90px;
  line-height:1.5;padding:2px 0;
}
#chat-input::placeholder{color:var(--t3)}
#voice-btn{
  width:32px;height:32px;border-radius:50%;flex-shrink:0;
  background:var(--s3);border:1px solid var(--border);
  display:flex;align-items:center;justify-content:center;
  cursor:pointer;transition:all .14s;
}
#voice-btn svg{width:13px;height:13px;fill:var(--t2);transition:fill .14s}
#voice-btn:active,#voice-btn.listening{
  background:rgba(255,68,56,.12);border-color:rgba(255,68,56,.3);
}
#voice-btn.listening{animation:pulse-mic .9s ease-in-out infinite}
#voice-btn.listening svg{fill:var(--err)}
@keyframes pulse-mic{0%,100%{box-shadow:0 0 0 0 rgba(255,68,56,.3)}50%{box-shadow:0 0 0 5px rgba(255,68,56,0)}}
#send-btn{
  width:32px;height:32px;border-radius:50%;flex-shrink:0;
  background:var(--acc);border:none;
  display:flex;align-items:center;justify-content:center;
  cursor:pointer;transition:all .12s;
}
#send-btn svg{width:13px;height:13px;fill:#000}
#send-btn:active{transform:scale(.9)}
#send-btn:disabled{opacity:.25;cursor:not-allowed;transform:none}

/* ════════════ TOAST ════════════ */
#toast{
  position:fixed;top:64px;left:50%;
  transform:translateX(-50%) translateY(-4px);
  background:var(--s4);border:1px solid var(--border2);
  color:var(--t1);padding:7px 14px;border-radius:20px;
  font-size:11px;font-weight:600;letter-spacing:.01em;
  opacity:0;transition:all .24s;z-index:300;
  pointer-events:none;white-space:nowrap;
  box-shadow:0 8px 24px rgba(0,0,0,.6);
}
#toast.show{opacity:1;transform:translateX(-50%) translateY(0)}

/* ════════════ DESKTOP ════════════ */
@media(min-width:500px){
  body{background:#050505;align-items:center}
  #hdr{max-width:420px;left:50%;transform:translateX(-50%);width:100%;
    border-left:1px solid var(--border);border-right:1px solid var(--border)}
  #page,#inp-wrap{max-width:420px;margin:0 auto}
}
</style>
</head>
<body>
<div id="toast"></div>

<!-- HEADER -->
<div id="hdr">
  <div id="hdr-av">SG</div>
  <div id="hdr-name">{{CREATOR_NAME}}</div>
  <div id="hdr-tag"><div class="live-dot"></div>KnowFlow</div>
</div>

<!-- PAGE -->
<div id="page">

  <!-- GREETING VIDEO -->
  <div id="video-sec">
    <div style="position:relative;width:100%;aspect-ratio:16/9;border-radius:12px;overflow:hidden;background:#000">
      <iframe id="video-player" src="" frameborder="0" allowfullscreen allow="autoplay; encrypted-media"
        style="position:absolute;inset:0;width:100%;height:100%;border:none"></iframe>
    </div>
  </div>

  <!-- PROFILE -->
  <div id="profile-sec">
    <div id="bio-text">{{CREATOR_BIO}}</div>
    <div id="tags-row"></div>
  </div>

  <!-- ACTION BUTTONS (links + products + booking unified) -->
  <div class="sec-label" id="action-label">Links</div>
  <div id="links-sec">
    <div class="act-grid" id="act-grid">
      <!-- filled by JS -->
    </div>
  </div>

  <div class="div"></div>

  <!-- XP -->
  <div id="xp-sec">
    <div class="xp-pill" id="lv-badge">Level 1</div>
    <div class="xp-pill streak-pill" id="streak-badge">0 Streak</div>
    <div id="xp-label">0 / 100 XP</div>
  </div>
  <div id="xp-track"><div id="xp-fill"></div></div>

  <!-- AI -->
  <div id="ai-sec">
    <div id="ai-hdr">
      <div id="ai-icon">
        <svg viewBox="0 0 24 24"><path d="M12 2a2 2 0 012 2c0 .74-.4 1.39-1 1.73V7h1a7 7 0 017 7h1v2h-1v1a2 2 0 01-2 2H5a2 2 0 01-2-2v-1H2v-2h1a7 7 0 017-7h1V5.73C10.4 5.39 10 4.74 10 4a2 2 0 012-2M7.5 13a1.5 1.5 0 00-1.5 1.5A1.5 1.5 0 007.5 16 1.5 1.5 0 009 14.5 1.5 1.5 0 007.5 13m9 0a1.5 1.5 0 00-1.5 1.5 1.5 1.5 0 001.5 1.5 1.5 1.5 0 001.5-1.5A1.5 1.5 0 0016.5 13z"/></svg>
      </div>
      <div>
        <div id="ai-label">AI Assistant</div>
        <div id="ai-sublabel">Trainiert auf {{CREATOR_NAME}}'s Content</div>
      </div>
    </div>

    <div id="qa-row">
      <div class="qa" data-qa="topics">
        <svg viewBox="0 0 24 24" width="11" height="11" fill="currentColor" style="flex-shrink:0"><path d="M4 6h16v2H4zm0 5h16v2H4zm0 5h16v2H4z"/></svg>
        Themen
      </div>
      <div class="qa" data-qa="start">
        <svg viewBox="0 0 24 24" width="11" height="11" fill="currentColor" style="flex-shrink:0"><path d="M12 17.27L18.18 21l-1.64-7.03L22 9.24l-7.19-.61L12 2 9.19 8.63 2 9.24l5.46 4.73L5.82 21z"/></svg>
        Start
      </div>
      <div class="qa" data-qa="goal">
        <svg viewBox="0 0 24 24" width="11" height="11" fill="currentColor" style="flex-shrink:0"><path d="M12 8a4 4 0 014 4 4 4 0 01-4 4 4 4 0 01-4-4 4 4 0 014-4m0-6a10 10 0 0110 10 10 10 0 01-10 10A10 10 0 012 12 10 10 0 0112 2z"/></svg>
        Mein Ziel
      </div>
      <div class="qa" data-qa="quiz">
        <svg viewBox="0 0 24 24" width="11" height="11" fill="currentColor" style="flex-shrink:0"><path d="M11.5 2C6.81 2 3 5.81 3 10.5S6.81 19 11.5 19h.5v3c4.86-2.34 8-7 8-11.5C20 5.81 16.19 2 11.5 2zm1 14.5h-2v-2h2v2zm0-4h-2c0-3.25 3-3 3-5 0-1.1-.9-2-2-2s-2 .9-2 2h-2c0-2.21 1.79-4 4-4s4 1.79 4 4c0 2.5-3 2.75-3 5z"/></svg>
        Quiz
      </div>
    </div>

    <div id="msg-list">
      <div class="msg msg-a">
        <div class="bbl">Hey — ich bin die AI von <strong>{{CREATOR_NAME}}</strong>, trainiert auf seinem gesamten Content.<br><br>Stell mir Fragen, teste dein Wissen oder sag mir was du erreichen willst.</div>
        <div class="msg-lbl">KnowFlow AI</div>
      </div>
    </div>
  </div>

  <a href="https://knowflow.io" target="_blank" id="pwrd">
    <div id="pwrd-dot"></div>Powered by KnowFlow
  </a>
  <a href="/admin?slug={{SLUG}}" style="display:block;text-align:center;font-size:9px;color:rgba(255,255,255,.18);text-decoration:none;margin-top:4px;margin-bottom:8px;font-family:DM Mono,monospace">✏️ admin</a>

</div><!-- /page -->

<!-- INPUT -->
<div id="inp-wrap">
  <div id="inp-actions">
    <div class="ia active-quiz" data-ia="quiz">
      <svg viewBox="0 0 24 24"><path d="M11.5 2C6.81 2 3 5.81 3 10.5S6.81 19 11.5 19h.5v3c4.86-2.34 8-7 8-11.5C20 5.81 16.19 2 11.5 2zm1 14.5h-2v-2h2v2zm0-4h-2c0-3.25 3-3 3-5 0-1.1-.9-2-2-2s-2 .9-2 2h-2c0-2.21 1.79-4 4-4s4 1.79 4 4c0 2.5-3 2.75-3 5z"/></svg>
      Quiz
    </div>
    <div class="ia" data-ia="topics">
      <svg viewBox="0 0 24 24"><path d="M4 6h16v2H4zm0 5h16v2H4zm0 5h16v2H4z"/></svg>
      Themen
    </div>
    <div class="ia" data-ia="start">
      <svg viewBox="0 0 24 24"><path d="M12 17.27L18.18 21l-1.64-7.03L22 9.24l-7.19-.61L12 2 9.19 8.63 2 9.24l5.46 4.73L5.82 21z"/></svg>
      Start
    </div>
    <div class="ia" data-ia="goal">
      <svg viewBox="0 0 24 24"><path d="M12 8a4 4 0 014 4 4 4 0 01-4 4 4 4 0 01-4-4 4 4 0 014-4m0-6a10 10 0 0110 10 10 10 0 01-10 10A10 10 0 012 12 10 10 0 0112 2z"/></svg>
      Ziel
    </div>
  </div>
  <div id="inp-row">
    <div id="voice-btn">
      <svg viewBox="0 0 24 24"><path d="M12 14c1.66 0 3-1.34 3-3V5c0-1.66-1.34-3-3-3S9 3.34 9 5v6c0 1.66 1.34 3 3 3zm5.91-3c-.49 0-.9.36-.98.85C16.52 14.2 14.47 16 12 16s-4.52-1.8-4.93-4.15c-.08-.49-.49-.85-.98-.85-.61 0-1.09.54-1 1.14.49 3 2.89 5.35 5.91 5.78V20c0 .55.45 1 1 1s1-.45 1-1v-2.08c3.02-.43 5.42-2.78 5.91-5.78.1-.6-.39-1.14-1-1.14z"/></svg>
    </div>
    <textarea id="chat-input" rows="1"
      placeholder="Frag {{CREATOR_NAME}}'s AI..."></textarea>
    <button id="send-btn" disabled>
      <svg viewBox="0 0 24 24"><path d="M2.01 21L23 12 2.01 3 2 10l15 2-15 2z"/></svg>
    </button>
  </div>
</div>

<script>
(function(){
'use strict';

const SLUG = '{{SLUG}}';
const BOOKING = '{{BOOKING_LINK}}';
const SID = 'sess_' + Math.random().toString(36).slice(2,11);

// STATE
let history = [], busy = false;
let xp = 0, level = 1, streak = 0;
let quizLog = []; // {ok, diff}
let voiceOn = false, recog = null;

// ── MARKDOWN ──
if (typeof marked !== 'undefined') marked.setOptions({breaks:true,gfm:true});
const md = t => {
  // Remove ** and * formatting using split (no regex needed)
  const clean = t.split('**').join('').split('__').join('');
  return typeof marked !== 'undefined' ? marked.parse(clean) : clean.split('\n').join('<br>');
};

// ── HELPERS ──
const $ = id => document.getElementById(id);
function toast(msg, ms=2200) {
  const el = $('toast');
  el.textContent = msg; el.classList.add('show');
  setTimeout(() => el.classList.remove('show'), ms);
}

// ── XP / LEVEL / STREAK ──
function gainXP(n) {
  xp += n;
  const cap = level * 100;
  if (xp >= cap) { xp -= cap; level++; toast('Level ' + level + ' erreicht', 3000); }
  $('xp-fill').style.width = Math.min(100, xp/(level*100)*100) + '%';
  $('xp-label').textContent = xp + ' / ' + (level*100) + ' XP';
  $('lv-badge').textContent = 'Level ' + level;
}
function setStreak(up) {
  if (up) { streak++; }
  else { streak = 0; }
  const el = $('streak-badge');
  el.textContent = streak + ' Streak';
  el.classList.toggle('hot', streak >= 3);
  if (streak >= 3) toast(streak + 'x Streak!');
}

// ── CONFIG LOAD ──
function getEmbedUrl(url) {
  if (!url) return null;
  url = url.trim();
  let yt = url.match(/(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([\w-]{11})/);
  if (yt) return 'https://www.youtube.com/embed/' + yt[1] + '?rel=0&modestbranding=1';
  let yts = url.match(/youtube\.com\/shorts\/([\w-]{11})/);
  if (yts) return 'https://www.youtube.com/embed/' + yts[1];
  let vm = url.match(/vimeo\.com\/(?:video\/)?([0-9]+)/);
  if (vm) return 'https://player.vimeo.com/video/' + vm[1] + '?color=ff6830&title=0&byline=0';
  return null;
}

async function loadConfig() {
  try {
    const cfg = await fetch('/api/creator/config?slug=' + SLUG).then(r => r.json());

    // greeting video - YouTube/Vimeo embed URL
    const videoUrl = cfg.greeting_video_url || '';
    if (videoUrl) {
      const embedUrl = getEmbedUrl(videoUrl);
      if (embedUrl) {
        $('video-player').src = embedUrl;
        $('video-sec').style.display = 'block';
      } else {
        $('video-sec').style.display = 'none';
      }
    } else {
      $('video-sec').style.display = 'none';
    }

    // bio + profile section
    const hasBio = cfg.bio && cfg.bio.trim();
    if (hasBio) {
      $('bio-text').textContent = cfg.bio.trim();
      $('profile-sec').style.display = '';
    } else {
      $('profile-sec').style.display = 'none';
    }

    // tags
    const tr = $('tags-row');
    if (cfg.video_count) tr.innerHTML += '<div class="tag accent">' + cfg.video_count + ' Videos</div>';
    tr.innerHTML += '<div class="tag">YouTube</div><div class="tag accent">AI aktiv</div>';

    // ── BUILD UNIFIED ACTION GRID ──
    // Collect all buttons: YouTube, products, booking
    const buttons = [];

    if (cfg.channel_url) {
      buttons.push({
        href: cfg.channel_url,
        cls: '',
        icon: '<svg viewBox="0 0 24 24"><path d="M10 15l5.19-3L10 9v6m11.56-7.83c.13.47.22 1.1.28 1.9.07.8.1 1.49.1 2.09L22 12c0 1.96-.14 3.41-.44 4.37-.27.84-.9 1.47-1.74 1.74-.47.13-1.33.22-2.65.28-1.3.07-2.49.1-3.59.1L12 18.5c-3.78 0-6.14-.14-7.06-.44a2.503 2.503 0 01-1.74-1.74C3 15.37 3 14.32 3 12c0-1.96.14-3.41.44-4.37.27-.84.9-1.47 1.74-1.74.47-.13 1.33-.22 2.65-.28 1.3-.07 2.49-.1 3.59-.1L12 5.5c3.78 0 6.14.14 7.06.44.84.27 1.47.9 1.74 1.74z"/></svg>',
        label: 'YouTube',
        sub: (cfg.video_count||0)+' Videos'
      });
    }

    if (cfg.products) {
      cfg.products.filter(p=>p.name).forEach(p => {
        buttons.push({
          href: p.url||'#',
          cls: '',
          icon: '<svg viewBox="0 0 24 24"><path d="M12 2l3.09 6.26L22 9.27l-5 4.87 1.18 6.88L12 17.77l-6.18 3.25L7 14.14 2 9.27l6.91-1.01L12 2z"/></svg>',
          label: p.name,
          sub: 'Angebot'
        });
      });
    }

    if (cfg.booking_link) {
      buttons.push({
        href: cfg.booking_link,
        cls: 'booking',
        icon: '<svg viewBox="0 0 24 24"><path d="M19 4h-1V2h-2v2H8V2H6v2H5C3.9 4 3 4.9 3 6v14c0 1.1.9 2 2 2h14c1.1 0 2-.9 2-2V6c0-1.1-.9-2-2-2zm0 16H5V9h14v11zM7 11h5v5H7z"/></svg>',
        label: 'Strategy Call',
        sub: 'Kostenlos buchen'
      });
    }

    const grid = $('act-grid');
    const count = buttons.length;
    // choose columns: 1→1col, 2→2col, 3→3col, 4→2x2, 5+→wrap
    if (count <= 3) grid.style.gridTemplateColumns = 'repeat('+count+',1fr)';
    else grid.style.gridTemplateColumns = 'repeat(3,1fr)';

    grid.innerHTML = buttons.map(b =>
      '<a href="'+b.href+'" target="_blank" class="act-btn '+b.cls+'">'
      + '<div class="act-icon">'+b.icon+'</div>'
      + '<div>'
      + '<div class="act-label">'+b.label+'</div>'
      + '<div class="act-sub">'+b.sub+'</div>'
      + '</div>'
      + '</a>'
    ).join('');

    if (!buttons.length) {
      $('action-label').style.display='none';
      $('links-sec').style.display='none';
    } else {
      $('action-label').style.display='';
      $('links-sec').style.display='';
    }

    // Update header name from config (in case template wasn't replaced)
    if (cfg.channel_name) {
      const nameEl = $('hdr-name');
      if (nameEl) nameEl.textContent = cfg.channel_name;
    }
    // creator initials in avatar
    const name = cfg.channel_name || '{{CREATOR_NAME}}';
    const parts = name.trim().split(' ');
    $('hdr-av').textContent = parts.length >= 2
      ? parts[0][0].toUpperCase() + parts[1][0].toUpperCase()
      : name.slice(0,2).toUpperCase();

  } catch(e) { console.warn('Config load error', e); }
}

// ── SCROLL ──
function scrollDown(el) {
  const p = $('page');
  if (el) {
    // Scroll so the element is visible with some padding
    setTimeout(() => {
      const rect = el.getBoundingClientRect();
      const pageRect = p.getBoundingClientRect();
      if (rect.bottom > pageRect.bottom - 60) {
        p.scrollTo({top: p.scrollTop + rect.bottom - pageRect.bottom + 80, behavior:'smooth'});
      }
    }, 80);
  } else {
    setTimeout(() => p.scrollTo({top: p.scrollHeight, behavior:'smooth'}), 60);
  }
}

// ── MESSAGES ──
function addMsg(role, content) {
  const list = $('msg-list');
  const wrap = document.createElement('div');
  wrap.className = 'msg msg-' + role;
  if (role === 'u') {
    wrap.innerHTML = '<div class="bbl">' + content.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/\n/g,'<br>') + '</div><div class="msg-lbl">Du</div>';
  } else {
    wrap.innerHTML = '<div class="bbl">' + md(content) + '</div><div class="msg-lbl">KnowFlow AI</div>';
  }
  list.appendChild(wrap);
  scrollDown(wrap);
}

function addOptions(opts) {
  const list = $('msg-list');
  const wrap = document.createElement('div');
  wrap.className = 'msg msg-a';
  wrap.innerHTML = '<div class="ans-opts">'
    + opts.map(o => '<button class="ans-o"><div class="ans-bullet"></div>' + o + '</button>').join('')
    + '</div>';
  wrap.querySelectorAll('.ans-o').forEach(btn => {
    btn.addEventListener('click', () => { wrap.remove(); sendMsg(btn.textContent.trim()); });
  });
  list.appendChild(wrap);
  scrollDown(wrap);
}

function showTyping() {
  const list = $('msg-list');
  const d = document.createElement('div');
  d.id = 'typing'; d.className = 'msg msg-a';
  d.innerHTML = '<div class="typing-bbl"><span></span><span></span><span></span></div>';
  list.appendChild(d); scrollDown();
}
function hideTyping() { const el = $('typing'); if (el) el.remove(); }

function renderBookingLinks(text) {
  let t = text;
  // Remove booking URLs - split on whitespace boundaries
  ['calendly.com/', 'cal.com/', 'tidycal.com/'].forEach(domain => {
    const parts = t.split('https://');
    t = parts.map((p, i) => {
      if (i === 0) return p;
      if (p.startsWith(domain) || p.startsWith('www.' + domain)) {
        // remove until whitespace or end
        return p.replace(/^[^\s\n)>\]"]+/, '');
      }
      return 'https://' + p;
    }).join('');
  });
  // Remove http:// variants too
  ['calendly.com/', 'cal.com/', 'tidycal.com/'].forEach(domain => {
    const parts = t.split('http://');
    t = parts.map((p, i) => {
      if (i === 0) return p;
      if (p.startsWith(domain)) return p.replace(/^[^\s\n)>\]"]+/, '');
      return 'http://' + p;
    }).join('');
  });
  return t.trim();
}

function showWarm() {
  const list = $('msg-list');
  const d = document.createElement('div');
  d.innerHTML = '<div class="warm-card">'
    + '<div class="warm-card-title">Bereit für den nächsten Schritt?</div>'
    + '<div class="warm-card-body">Basierend auf unserem Gespräch könnte ein persönlicher Strategy Call genau das Richtige sein.</div>'
    + '<a href="'+(BOOKING||'#')+'" target="_blank" class="warm-btn">Call buchen'
    + '<svg viewBox="0 0 24 24"><path d="M8.59 16.59L13.17 12 8.59 7.41 10 6l6 6-6 6z"/></svg></a>'
    + '</div>';
  $('msg-list').appendChild(d); scrollDown(); gainXP(20);
}

// ── QUIZ ──
function nextDiff() {
  if (quizLog.length < 2) return 1;
  const r = quizLog.slice(-3);
  const ok = r.filter(x=>x.ok).length;
  const last = r[r.length-1].diff;
  if (ok === r.length && last < 3) return Math.min(3, last+1);
  if (ok === 0 && last > 1) return Math.max(1, last-1);
  return last;
}

const diffLabel = {1:'Leicht',2:'Mittel',3:'Schwer'};

function renderQuiz(quiz, diff) {
  const list = $('msg-list');
  const fbId = 'fb'+Date.now(), expId = 'exp'+Date.now();
  const opts = quiz.options.map((o, i) => {
    const k = 'ABCD'[i];
    const lbl = o.replace(/^[A-D][)\.]\s*/,'');
    return '<button class="qz-o" data-k="'+k+'" data-correct="'+quiz.correct+'" data-expid="'+expId+'" data-diff="'+diff+'">'
      + '<span class="qz-k">'+k+'</span>' + lbl + '</button>';
  }).join('');
  const card = document.createElement('div');
  card.className = 'qz';
  card.innerHTML = '<div class="qz-tag">Quiz<span class="qz-diff-tag">'+diffLabel[diff]+'</span></div>'
    + '<div class="qz-q">'+quiz.question+'</div>'
    + '<div class="qz-opts">'+opts+'</div>'
    + '<div class="qz-fb" id="'+fbId+'"></div>'
    + '<div class="qz-exp" id="'+expId+'"></div>';

  card.querySelectorAll('.qz-o').forEach(btn => btn.addEventListener('click', () => answerQuiz(btn, fbId, expId)));
  list.appendChild(card); scrollDown();
}

function answerQuiz(btn, fbId, expId) {
  const card = btn.closest('.qz');
  const correct = btn.dataset.correct;
  const diff = parseInt(btn.dataset.diff);
  const sel = btn.dataset.k;
  card.querySelectorAll('.qz-o').forEach(b => {
    b.disabled = true;
    if (b.dataset.k === correct) b.classList.add('correct');
    else if (b.dataset.k === sel && sel !== correct) b.classList.add('wrong');
  });
  const ok = sel === correct;
  quizLog.push({ok, diff});
  const fb = $(fbId);
  fb.className = 'qz-fb show ' + (ok ? 'ok' : 'fail');
  if (ok) {
    gainXP(15); setStreak(true);
    fb.innerHTML = '<span class="qz-fb-icon">✓</span> Richtig — +15 XP';
    // Duolingo-style: next question auto after short pause
    setTimeout(() => triggerQuiz(), 1200);
  } else {
    setStreak(false);
    fb.innerHTML = '<span class="qz-fb-icon">✗</span> Richtig wäre: <strong>' + correct + '</strong>';
    // Fetch explanation, then auto-next after reading time
    fetchExplanation(expId, card.querySelector('.qz-q').textContent, correct, () => {
      setTimeout(() => triggerQuiz(), 2800);
    });
  }
}

async function fetchExplanation(expId, question, correct, onDone) {
  try {
    const r = await fetch('/api/quiz-explain', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({slug:SLUG, question, correct_answer:correct})
    });
    const d = await r.json();
    if (d.explanation) {
      const el = $(expId);
      if (el) { el.textContent = d.explanation; el.classList.add('show'); scrollDown(); }
    }
  } catch(e) {}
  if (onDone) onDone();
}

async function triggerTopics() {
  if (busy) return;
  showTyping();
  busy = true;
  try {
    const r = await fetch('/api/chat', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({
        slug: SLUG,
        messages: [{role:'user', content: 'Liste mir die 6 wichtigsten Themen auf über die du mich unterrichten kannst. Antworte NUR mit einer nummerierten Liste, jedes Thema max 5 Wörter, keine Erklärungen.'}],
        session_id: SID
      })
    });
    const d = await r.json();
    hideTyping();
    busy = false;
    if (d.response) {
      // parse topics from numbered list
      const lines = d.response.split('\n');
      const topics = [];
      for (const l of lines) {
        const m = l.match(/^\d+[.)\s]+(.{3,60})$/);
        if (m) topics.push(m[1].trim().replace(/[*_]/g,''));
      }
      if (topics.length >= 2) {
        addTopicChips(topics);
      } else {
        addMsg('a', d.response);
      }
    }
  } catch(e) { hideTyping(); busy = false; }
}

function addTopicChips(topics) {
  busy = false; // always reset after topics loaded
  $('send-btn').disabled = false;
  const list = $('msg-list');
  const lbl = document.createElement('div');
  lbl.className = 'msg msg-a';
  lbl.innerHTML = '<div class="bbl"><span style="color:var(--t2);font-size:12px">Wähle ein Thema:</span></div><div class="msg-lbl">KnowFlow AI</div>';
  list.appendChild(lbl);
  // chip container
  const wrap = document.createElement('div');
  wrap.className = 'topic-chips';
  topics.forEach(t => {
    const btn = document.createElement('button');
    btn.className = 'topic-chip';
    btn.textContent = t;
    btn.addEventListener('click', () => {
      wrap.remove();
      sendMsg('Erkläre mir das Thema: ' + t);
    });
    wrap.appendChild(btn);
  });
  list.appendChild(wrap);
  scrollDown();
}

async function triggerQuiz() {
  const diff = nextDiff();
  showTyping();
  try {
    const r = await fetch('/api/quiz', {
      method:'POST', headers:{'Content-Type':'application/json'},
      body: JSON.stringify({slug:SLUG, difficulty:diff})
    });
    const d = await r.json();
    hideTyping();
    if (d.quiz) renderQuiz(d.quiz, diff);
  } catch(e) { hideTyping(); }
}

// ── VOICE ──
// ── VOICE with WhatsApp-style waveform ──
let waveAnim = null, waveSec = 0;

function showVoiceWave() {
  $('chat-input').style.display = 'none';
  $('send-btn').style.display = 'none';
  const wave = document.createElement('div');
  wave.id = 'voice-wave';
  wave.innerHTML =
    '<button id="voice-cancel" title="Abbrechen"><svg viewBox="0 0 24 24" width="14" height="14" fill="currentColor"><path d="M19 6.41L17.59 5 12 10.59 6.41 5 5 6.41 10.59 12 5 17.59 6.41 19 12 13.41 17.59 19 19 17.59 13.41 12z"/></svg></button>'
    + '<div class="wave-bars">' + Array.from({length:24},(_,i)=>'<div class="wave-bar" id="wb'+i+'"></div>').join('') + '</div>'
    + '<span id="voice-dur">0:00</span>';
  $('inp-row').appendChild(wave);
  waveSec = 0;
  waveAnim = setInterval(() => {
    for (let i=0;i<24;i++){const b=$('wb'+i);if(b)b.style.height=(3+Math.abs(Math.sin(Date.now()/200+i*.5))*18+Math.random()*6)+'px';}
    waveSec++;
    const d=$('voice-dur'); if(d)d.textContent=Math.floor(waveSec/10)+':'+(String(waveSec%10*6).padStart(2,'0'));
  }, 100);
  $('voice-cancel').addEventListener('click', ()=>{ if(recog)recog.abort(); stopVoiceWave(true); });
}

function stopVoiceWave(cancelled) {
  clearInterval(waveAnim); waveAnim = null; waveSec = 0;
  voiceOn = false;
  $('voice-btn').classList.remove('listening');
  const wave = $('voice-wave'); if(wave) wave.remove();
  $('chat-input').style.display = '';
  $('send-btn').style.display = '';
  updateSendBtn();
}

function toggleVoice() {
  const SR = window.SpeechRecognition || window.webkitSpeechRecognition;
  if (!SR) { alert('Spracherkennung nicht verfügbar in diesem Browser. Bitte Chrome oder Safari nutzen.'); return; }
  const btn = $('voice-btn');
  if (voiceOn) { if(recog) recog.stop(); stopVoiceWave(); return; }
  recog = new SR();
  recog.lang = 'de-DE';
  recog.interimResults = true;   // get partial results too
  recog.continuous = false;
  recog.maxAlternatives = 1;
  let finalTranscript = '';
  recog.onstart = () => { voiceOn = true; btn.classList.add('listening'); showVoiceWave(); finalTranscript = ''; };
  recog.onresult = e => {
    let interim = '';
    for (let i = e.resultIndex; i < e.results.length; i++) {
      const t = e.results[i][0].transcript;
      if (e.results[i].isFinal) finalTranscript += t;
      else interim = t;
    }
    // Show interim in input as preview
    $('chat-input').value = finalTranscript || interim;
    updateSendBtn();
  };
  recog.onend = () => {
    const text = finalTranscript || $('chat-input').value.trim();
    stopVoiceWave();
    if (text) {
      $('chat-input').value = text;
      updateSendBtn();
      setTimeout(() => sendMsg(), 100);
    }
  };
  recog.onerror = e => {
    console.warn('Voice error:', e.error);
    stopVoiceWave(true);
    if (e.error === 'not-allowed') toast('Mikrofon-Zugriff verweigert');
    else if (e.error === 'no-speech') toast('Kein Ton erkannt — nochmal versuchen');
  };
  recog.start();
}

// ── QUICK ACTIONS ──
const qaMap = {
  topics: '__TOPICS__', // special: fetch topics and show as clickable chips
  start: 'Was sollte ich als Erstes lernen und warum?',
  goal: 'Ich möchte dir erklären was ich erreichen möchte',
  quiz: null // special
};

// ── SEND ──
function sendMsg(text) {
  const inp = $('chat-input');
  const msg = (text || inp.value).trim();
  if (!msg || busy) return;
  inp.value = ''; inp.style.height = 'auto';
  updateSendBtn();
  busy = true; $('send-btn').disabled = true;
  addMsg('u', msg); gainXP(5);
  history.push({role:'user', content:msg});
  showTyping();
  fetch('/api/chat', {
    method:'POST', headers:{'Content-Type':'application/json'},
    body: JSON.stringify({slug:SLUG, messages:history, session_id:SID})
  })
  .then(r => r.json())
  .then(d => {
    hideTyping();
    if (d.response) {
      const cleanedResponse = renderBookingLinks(d.response);
      addMsg('a', cleanedResponse);
      history.push({role:'assistant', content:d.response});
      gainXP(5);
      // Show booking card if warm_lead OR if response contains a booking URL
      const hasBookingUrl = BOOKING && d.response.toLowerCase().includes(BOOKING.toLowerCase().slice(0,20));
      const hasCalendly = d.response.includes('calendly.com') || d.response.includes('cal.com') || d.response.includes('tidycal.com');
      if ((d.warm_lead || hasBookingUrl || hasCalendly) && BOOKING && !document.querySelector('.warm-card')) {
        setTimeout(() => showWarm(), 400);
      }
      const opts = extractOptions(d.response);
      if (opts.length >= 2) addOptions(opts);
    } else {
      const errMsg = d.error || 'Verbindungsfehler. Nochmal versuchen.';
      addMsg('a', errMsg);
    }
  })
  .catch(() => { hideTyping(); addMsg('a', 'Verbindungsfehler. Nochmal versuchen.'); })
  .finally(() => { busy = false; $('send-btn').disabled = false; inp.focus(); });
}

function extractOptions(text) {
  const lines = text.split('\n');
  const opts = [];
  for (const l of lines) {
    const m = l.match(/^\d+[.)]\s+(.{4,100})$/);
    if (m) {
      const clean = m[1].trim().replace(/[*][*]/g,'').replace(/[*]/g,'');
      opts.push(clean);
    }
  }
  return opts.length >= 2 && opts.length <= 6 ? opts : [];
}

// ── INPUT EVENTS ──
const inp = $('chat-input');
function updateSendBtn() {
  $('send-btn').disabled = inp.value.trim().length === 0 || busy;
}
inp.addEventListener('input', () => {
  inp.style.height = 'auto';
  inp.style.height = Math.min(inp.scrollHeight, 90) + 'px';
  updateSendBtn();
});
inp.addEventListener('keydown', e => {
  if (e.key === 'Enter' && !e.shiftKey) { e.preventDefault(); sendMsg(); }
});

// QA chips (both in header and input bar)
document.querySelectorAll('[data-qa],[data-ia]').forEach(el => {
  const key = el.dataset.qa || el.dataset.ia;
  el.addEventListener('click', () => {
    if (key === 'quiz') triggerQuiz();
    else if (key === 'topics') triggerTopics();
    else if (qaMap[key]) sendMsg(qaMap[key]);
  });
});

$('send-btn').addEventListener('click', () => sendMsg());
$('voice-btn').addEventListener('click', toggleVoice);

loadConfig();

})();
</script>
</body>
</html>"""


# ── Main ──────────────────────────────────────────────────────────────────────
def main():
    port = int(os.environ.get("PORT", 7891))
    print(f"""
╔══════════════════════════════════════════════════════╗
║   KnowFlow MVP v1.0                                  ║
║   YouTube Knowledge → AI Learning + Sales Funnel     ║
╚══════════════════════════════════════════════════════╝

  🎬 Creator Setup:  http://localhost:{port}/
  📖 Viewer Page:    http://localhost:{port}/c/YOUR-SLUG

  Requirements:
  → Set ANTHROPIC_API_KEY environment variable
    export ANTHROPIC_API_KEY=sk-ant-...

""")
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key or not api_key.startswith("sk-"):
        print("⚠️  WARNUNG: ANTHROPIC_API_KEY nicht gesetzt!")
        print("   Setze ihn mit: export ANTHROPIC_API_KEY=sk-ant-...")
        print("   Dann Server neu starten.\n")
    else:
        print(f"  ✅ API Key: ...{api_key[-8:]}\n")
    if not os.environ.get("PORT"):  # only open browser locally
        threading.Timer(1.2, lambda: webbrowser.open(f"http://localhost:{port}")).start()
    try:
        HTTPServer(("0.0.0.0", port), Handler).serve_forever()
    except KeyboardInterrupt:
        print("\n👋 KnowFlow stopped.")

if __name__ == "__main__":
    main()
